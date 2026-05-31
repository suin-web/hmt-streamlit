# -*- coding: utf-8 -*-
"""
학생맞춤통합지원 AI 교사용 의사결정 보조 프로토타입

실행:
    streamlit run hakmatch_streamlit_app.py

이번 버전 반영 사항:
- 접속 첫 화면: 교사 대시보드
- 상단 역할 전환: 담임교사 / 학생맞춤통합지원담당교원
- 담임교사: 본인 반 학생만 조회
- 학생맞춤통합지원담당교원: 전교 학생 상태 조회
- 학교 정보 자료 연결
- 1차 체크리스트 입력 → 점수 계산 → 맥락 보정 → 우선 확인 신호 확인
  → 심층 유도 분석 활성화 → 상담지 생성 고려 영역 산출
- Gemini API 기반 2차 상담 질문 생성, 상담 메모 구조화, 기관 추천 이유 생성
- 검색 자료 기반 맞춤 검색, docx 템플릿 기반 회의록 생성

주의:
- 이 코드는 발표·시연용 MVP입니다.
- 실제 학생 개인정보, 실제 학교명, 민감 정보는 넣지 마세요.
- 문서 생성 단계에서 개인정보는 LLM에 보내지 않고 Python이 템플릿에 직접 삽입합니다.
"""

from __future__ import annotations

import io
import copy
import json
import html
import hashlib
import math
import os
import re
from datetime import date, datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

import pandas as pd
import streamlit as st

from gemini_client import call_llm_with_validation, get_gemini_api_key
from validation import (
    HARD_BANNED_COMMON,
    validate_counseling_question_output,
    validate_counseling_analysis_output,
    validate_resource_recommendation_output,
    validate_document_generation_output,
    sanitize_llm_parsed_data,
)

try:
    import plotly.express as px
except Exception:
    px = None

# -----------------------------------------------------------------------------
# 기본 설정
# -----------------------------------------------------------------------------
st.set_page_config(
    page_title="학생맞춤통합지원 AI | 교사용 지원 신호 대시보드",
    page_icon="🧭",
    layout="wide",
    initial_sidebar_state="expanded",
)

APP_DIR = Path(__file__).resolve().parent
DATA_DIR = APP_DIR / "data"

# 파일 경로는 여기만 바꾸면 됩니다.
CSV_PATHS = {
    "school_info": [
        DATA_DIR / "02_school_info_db_location.csv",
        DATA_DIR / "05_school_info_db_location.csv",
        APP_DIR / "02_school_info_db_location.csv",
        APP_DIR / "05_school_info_db_location.csv",
    ],
    "school_context": [
        DATA_DIR / "02_school_context_scores_db.csv",
        APP_DIR / "02_school_context_scores_db.csv",
    ],
    "region_context": [
        DATA_DIR / "서울_학생지원맥락_영역별점수.csv",
        APP_DIR / "서울_학생지원맥락_영역별점수.csv",
    ],
    "checklist": [
        DATA_DIR / "first_checklist_items_v1.csv",
        APP_DIR / "first_checklist_items_v1.csv",
    ],
    "deep_rules": [
        DATA_DIR / "deep_inference_rules_v1.csv",
        APP_DIR / "deep_inference_rules_v1.csv",
    ],
    "rule_map": [
        DATA_DIR / "checklist_item_deep_rule_map_v1.csv",
        APP_DIR / "checklist_item_deep_rule_map_v1.csv",
    ],
    "official_checklist": [
        DATA_DIR / "official_counseling_checklist_reference_v1.csv",
        APP_DIR / "official_counseling_checklist_reference_v1.csv",
    ],
}

JSON_PATHS = {
    "district_adjacency": [
        DATA_DIR / "district_adjacency_seoul_FILTER_READY_20260527_v3.json",
        DATA_DIR / "district_adjacency_seoul_v1.json",
        APP_DIR / "district_adjacency_seoul_FILTER_READY_20260527_v3.json",
        APP_DIR / "district_adjacency_seoul_v1.json",
    ],
    "rag_config": [
        DATA_DIR / "rag_filter_rank_config_20260527_v3.json",
        APP_DIR / "rag_filter_rank_config_20260527_v3.json",
    ],
}

TEMPLATE_DIR = APP_DIR / "templates"
OUTPUT_DIR = APP_DIR / "outputs"

# 기존 코드 변수명이 다를 때 이곳만 맞추면 됩니다.
SCHOOL_CONTEXT_VAR_NAME = "selected_school_context"
REGION_CONTEXT_VAR_NAME = "selected_region_context"

SUPPORT_AREAS = ["학업", "심리정서", "복지경제", "진로"]
AREA_TO_SCHOOL_CONTEXT_COL = {
    "학업": "학교_학습점수",
    "심리정서": "학교_정서심리점수",
    "복지경제": "학교_경제복지점수",
    "진로": "학교_진로점수",
}
AREA_TO_REGION_CONTEXT_COL = {
    "학업": "지역_학습점수",
    "심리정서": "지역_정서심리점수",
    "복지경제": "지역_경제복지점수",
    "진로": "지역_진로점수",
}

STATUS_ORDER = ["심층 파악 필요", "심층 파악 권고", "주의 및 탐색", "일상적 관찰"]
ROLE_HOMEROOM = "담임교사"
ROLE_COORDINATOR = "학생맞춤통합지원담당교원"
DEFAULT_GRADE = "2학년"
DEFAULT_CLASS = "3반"

# -----------------------------------------------------------------------------
# CSS
# -----------------------------------------------------------------------------
def inject_css() -> None:
    st.markdown(
        """
        <style>
        :root {
            --blue: #2f6bff;
            --deep-blue: #0b2a55;
            --navy: #111927;
            --line: #d8dee9;
            --muted: #64748b;
            --bg: #f4f7fb;
            --soft-blue: #eef4ff;
            --green: #14a36f;
            --orange: #f59e0b;
            --red: #ef4444;
            --purple: #7c3aed;
        }
        .stApp { background: var(--bg); }
        section[data-testid="stSidebar"] {
            background: #ffffff;
            border-right: 1px solid #d6dbe6;
        }
        section[data-testid="stSidebar"] .block-container { padding-top: 1rem; }
        .edutop {
            display: flex;
            align-items: stretch;
            justify-content: space-between;
            border: 1px solid #cbd5e1;
            background: #ffffff;
            margin: -0.5rem 0 0.8rem 0;
            min-height: 74px;
        }
        .edu-logo {
            width: 210px;
            display: flex;
            align-items: center;
            gap: 10px;
            padding: 12px 16px;
            border-right: 1px solid #cbd5e1;
            font-weight: 900;
            color: #0f172a;
        }
        .edu-symbol {
            width: 34px;
            height: 34px;
            border-radius: 50%;
            background: linear-gradient(135deg, #2f6bff, #00a884);
            display: inline-flex;
            align-items: center;
            justify-content: center;
            color: white;
            font-size: 18px;
        }
        .edu-main { flex: 1; }
        .edu-bluebar {
            height: 34px;
            background: var(--blue);
            color: white;
            display: flex;
            align-items: center;
            justify-content: space-between;
            padding: 0 12px;
            font-size: 0.88rem;
        }
        .edu-school-badge {
            background: rgba(255,255,255,0.18);
            border: 1px solid rgba(255,255,255,0.35);
            border-radius: 4px;
            padding: 3px 8px;
            font-weight: 800;
            margin-right: 8px;
        }
        .edu-actions span {
            border: 1px solid rgba(255,255,255,0.7);
            border-radius: 3px;
            padding: 3px 8px;
            margin-left: 5px;
            background: #ffffff;
            color: #0f172a;
            font-weight: 800;
        }
        .rolebar {
            min-height: 40px;
            display: flex;
            align-items: center;
            justify-content: space-between;
            gap: 12px;
            padding: 0 18px;
            font-weight: 800;
            color: #0f2a50;
            border-top: 1px solid #cbd5e1;
            background: #ffffff;
        }
        .role-pill {
            display: inline-block;
            border-radius: 999px;
            padding: 5px 10px;
            background: #eef4ff;
            border: 1px solid #bfdbfe;
            color: #1d4ed8;
            font-size: .82rem;
        }
        .page-title {
            font-size: 1.18rem;
            font-weight: 900;
            color: #0f172a;
            margin: 0.3rem 0 0.2rem 0;
        }
        .page-subtitle {
            font-size: 0.92rem;
            color: #64748b;
            margin-bottom: 0.8rem;
        }
        .panel {
            background: #ffffff;
            border: 1px solid #dbe2ef;
            border-radius: 10px;
            padding: 16px;
            box-shadow: 0 1px 2px rgba(15, 23, 42, 0.04);
            margin-bottom: 12px;
        }
        .panel-title {
            font-weight: 900;
            color: #0f2a50;
            margin-bottom: 10px;
            display: flex;
            align-items: center;
            gap: 6px;
        }
        .panel-title::before {
            content: '';
            display: inline-block;
            width: 5px;
            height: 18px;
            border-radius: 4px;
            background: var(--blue);
        }
        .metric-card {
            background: #ffffff;
            border: 1px solid #dbe2ef;
            border-radius: 12px;
            padding: 16px 16px 14px 16px;
            box-shadow: 0 1px 3px rgba(15,23,42,0.05);
            min-height: 96px;
        }
        .metric-label {
            font-size: 0.82rem;
            color: #64748b;
            font-weight: 800;
            margin-bottom: 8px;
        }
        .metric-value {
            font-size: 1.28rem;
            line-height: 1.25;
            word-break: keep-all;
            color: #0f172a;
            font-weight: 900;
        }
        .metric-help {
            font-size: 0.78rem;
            color: #64748b;
            margin-top: 8px;
            line-height: 1.45;
            word-break: keep-all;
        }
        .student-card {
            border: 1px solid #dbe2ef;
            background: #ffffff;
            border-radius: 12px;
            padding: 14px 16px;
            margin-bottom: 10px;
            box-shadow: 0 1px 2px rgba(15,23,42,0.04);
        }
        .risk-deep { border-left: 8px solid var(--red); background: #fff7f7; }
        .risk-watch { border-left: 8px solid var(--orange); background: #fffaf0; }
        .risk-normal { border-left: 8px solid #94a3b8; }
        .badge {
            display: inline-block;
            border-radius: 999px;
            padding: 3px 9px;
            font-weight: 900;
            font-size: 0.78rem;
            margin-right: 5px;
            border: 1px solid transparent;
        }
        .badge-deep { background: #fee2e2; color: #b91c1c; border-color: #fecaca; }
        .badge-watch { background: #ffedd5; color: #c2410c; border-color: #fed7aa; }
        .badge-normal { background: #f1f5f9; color: #475569; border-color: #e2e8f0; }
        .mini-grid {
            display: grid;
            grid-template-columns: repeat(4, minmax(80px, 1fr));
            gap: 7px;
            margin-top: 10px;
        }
        .mini-domain {
            background: #f8fafc;
            border: 1px solid #e2e8f0;
            border-radius: 8px;
            padding: 7px;
        }
        .domain-name { font-size: 0.72rem; color: #64748b; font-weight: 800; margin-bottom: 4px; }
        .dots { letter-spacing: 1px; color: #2563eb; font-weight: 900; }
        .dots-zero { color: #cbd5e1; }
        .info-table {
            width: 100%;
            border-collapse: collapse;
            font-size: 0.86rem;
        }
        .info-table th {
            background: #f1f5f9;
            color: #0f2a50;
            text-align: left;
            padding: 9px;
            border: 1px solid #dbe2ef;
        }
        .info-table td {
            padding: 9px;
            border: 1px solid #dbe2ef;
            vertical-align: top;
        }
        .small-muted { color: #64748b; font-size: 0.84rem; }
        .recommend-card {
            background: #ffffff;
            border: 1px solid #dbe2ef;
            border-radius: 12px;
            padding: 14px 16px;
            margin-bottom: 10px;
        }
        .recommend-rank {
            display: inline-flex;
            width: 25px;
            height: 25px;
            border-radius: 50%;
            align-items: center;
            justify-content: center;
            background: var(--deep-blue);
            color: white;
            font-weight: 900;
            margin-right: 8px;
        }
        .question-text {
            font-weight: 850;
            font-size: 1rem;
            line-height: 1.55;
            color: #111827;
            margin: 2px 0 10px 0;
        }
        .analysis-summary-card {
            background: #ffffff;
            border: 1px solid #dbe2ef;
            border-radius: 12px;
            padding: 18px 20px;
            min-height: 150px;
            box-shadow: 0 1px 3px rgba(15,23,42,0.05);
        }
        .analysis-summary-label {
            font-size: .9rem;
            color: #64748b;
            font-weight: 850;
            margin-bottom: 8px;
        }
        .analysis-summary-status {
            font-size: 1.42rem;
            line-height: 1.25;
            font-weight: 900;
            color: #0f172a;
            margin-bottom: 12px;
        }
        .analysis-summary-text {
            font-size: 1rem;
            line-height: 1.7;
            color: #334155;
            word-break: keep-all;
        }
        .resource-title {
            font-weight: 900;
            color: #0f172a;
            font-size: 1.05rem;
            margin: 4px 0 10px 0;
        }
        .resource-detail-table {
            width: 100%;
            border-collapse: collapse;
            font-size: .92rem;
            margin-bottom: 8px;
        }
        .resource-detail-table th {
            width: 150px;
            background: #f8fafc;
            color: #334155;
            text-align: left;
            padding: 9px 10px;
            border: 1px solid #dbe2ef;
        }
        .resource-detail-table td {
            padding: 9px 10px;
            border: 1px solid #dbe2ef;
            color: #0f172a;
            line-height: 1.55;
        }
        .callout {
            padding: 12px 14px;
            border-radius: 10px;
            border: 1px solid #bfdbfe;
            background: #eff6ff;
            color: #0f2a50;
            margin-bottom: 10px;
        }
        .warning-callout {
            padding: 12px 14px;
            border-radius: 10px;
            border: 1px solid #fecaca;
            background: #fff7f7;
            color: #7f1d1d;
            margin-bottom: 10px;
        }

        .workflow-section-header {
            background: #ffffff;
            border: 1px solid #dbe2ef;
            border-radius: 14px;
            padding: 18px 20px;
            margin: 22px 0 14px 0;
            box-shadow: 0 1px 3px rgba(15,23,42,0.04);
        }
        .workflow-section-kicker {
            display: inline-flex;
            align-items: center;
            justify-content: center;
            min-width: 32px;
            height: 26px;
            border-radius: 999px;
            background: #eef4ff;
            border: 1px solid #bfdbfe;
            color: #1d4ed8;
            font-weight: 900;
            font-size: .82rem;
            margin-right: 8px;
        }
        .workflow-section-title {
            font-size: 1.08rem;
            font-weight: 900;
            color: #0f172a;
            margin-bottom: 4px;
        }
        .workflow-section-subtitle {
            color: #64748b;
            font-size: .9rem;
            line-height: 1.5;
        }
        .area-chip-row {
            display: flex;
            flex-wrap: wrap;
            gap: 7px;
            margin-top: 10px;
        }
        .area-chip-main {
            display: inline-flex;
            align-items: center;
            border-radius: 999px;
            padding: 7px 12px;
            background: #eff6ff;
            border: 1px solid #bfdbfe;
            color: #1e40af;
            font-size: .95rem;
            font-weight: 900;
        }
        .area-chip-sub {
            display: inline-flex;
            align-items: center;
            border-radius: 999px;
            padding: 5px 10px;
            background: #f8fafc;
            border: 1px solid #e2e8f0;
            color: #64748b;
            font-size: .8rem;
            font-weight: 750;
        }
        .analysis-readable-text {
            font-size: 1.03rem;
            line-height: 1.85;
            color: #1f2937;
            word-break: keep-all;
            margin-top: 6px;
        }
        .footer-note {
            color: #64748b;
            font-size: 0.82rem;
            border-top: 1px solid #e2e8f0;
            padding-top: 10px;
            margin-top: 18px;
        }
        
.soft-card.compact-card {
    background:#ffffff;
    border:1px solid #d8e2ef;
    border-radius:16px;
    padding:18px 20px;
    margin:12px 0 12px 0;
    box-shadow:0 4px 12px rgba(15,23,42,.04);
}
.readable-note {
    font-size:1.02rem;
    line-height:1.75;
    color:#243044;
    margin-top:8px;
}
.area-inline {
    display:inline-block;
    margin-top:8px;
    padding:8px 12px;
    border-radius:999px;
    background:#eef5ff;
    color:#1d4ed8;
    font-weight:800;
}

</style>
        """,
        unsafe_allow_html=True,
    )

# -----------------------------------------------------------------------------
# 공통 로더
# -----------------------------------------------------------------------------
def first_existing_path(paths: Iterable[Path]) -> Optional[Path]:
    for path in paths:
        if path.exists():
            return path
    return None


def load_csv_with_fallback(path: Path | str) -> pd.DataFrame:
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(str(path))
    try:
        return pd.read_csv(path, encoding="utf-8-sig")
    except UnicodeDecodeError:
        return pd.read_csv(path, encoding="cp949")


def load_optional_csv(paths: Iterable[Path], name: str) -> Tuple[Optional[pd.DataFrame], Optional[Path], Optional[str]]:
    path = first_existing_path(paths)
    if path is None:
        return None, None, f"{name} 파일을 찾을 수 없습니다. 확인한 경로: " + ", ".join(str(p) for p in paths)
    try:
        return load_csv_with_fallback(path), path, None
    except Exception as e:
        return None, path, f"{path.name} 읽기 실패: {e}"


def get_cell(row: Optional[pd.Series | Dict[str, Any]], col: str, default: float = 0.0) -> float:
    if row is None:
        return default
    try:
        value = row[col]
    except Exception:
        return default
    try:
        if pd.isna(value):
            return default
    except Exception:
        pass
    try:
        return float(value)
    except Exception:
        return default


def normalize_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and math.isnan(value):
        return ""
    return str(value).strip()


def normalize_area_name(value: str) -> str:
    value = normalize_text(value)
    mapping = {
        "학습": "학업",
        "학습지원": "학업",
        "학습·진로": "학업",
        "학업진로": "학업",
        "정서": "심리정서",
        "심리": "심리정서",
        "심리·정서": "심리정서",
        "심리정서지원": "심리정서",
        "복지": "복지경제",
        "경제": "복지경제",
        "복지·경제": "복지경제",
        "환경": "복지경제",
        "건강·안전": "심리정서",
        "교사관계": "심리정서",
    }
    return mapping.get(value, value)




def display_area_name(value: Any) -> str:
    """내부 계산용 영역명을 사용자 화면용 용어로 바꾼다."""
    area = normalize_area_name(value)
    mapping = {
        "학업": "학습",
        "복지경제": "복지경제",
        "심리정서": "심리정서",
        "진로": "진로",
        "긴급확인": "긴급확인",
        "공통": "공통",
    }
    return mapping.get(area, normalize_text(value))


def display_area_list(values: Iterable[Any]) -> List[str]:
    result: List[str] = []
    for value in values or []:
        name = display_area_name(value)
        if name and name not in result:
            result.append(name)
    return result


def display_text(value: Any) -> str:
    """사용자에게 보이는 문구에서 내부 개발용 용어를 정리한다."""
    text = normalize_text(value)
    if not text:
        return ""
    replacements = {
        "복지 및 환경지원 영역": "복지 및 경제 지원 영역",
        "복지 및 환경 지원 영역": "복지 및 경제 지원 영역",
        "복지·환경": "복지·경제",
        "환경지원": "경제 지원",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    # 학습-A, 정서-C처럼 내부 분석용으로 붙인 문항 코드는 사용자 화면에서 제거한다.
    text = re.sub(r"(?:학습|정서|환경|진로)-[A-Z](?:·(?:학습|정서|환경|진로)-[A-Z])*\s*심층 유도[:：]?\s*", "", text)
    text = re.sub(r"(?:학습|정서|환경|진로)-[A-Z]", "", text)
    text = re.sub(r"\s{2,}", " ", text).strip(" .·-:")
    # 학업중단처럼 고유 표현은 유지하고, 영역명으로 쓰인 학업만 학습으로 표기한다.
    text = re.sub(r"(?<![가-힣])학업(?![가-힣])", "학습", text)
    text = text.replace("학업 영역", "학습 영역")
    text = text.replace("학업 지원", "학습 지원")
    return text


def clean_checklist_item_label(item_text: Any) -> str:
    """학습-A, 정서-C 같은 내부 문항 코드를 사용자 화면에서 제거한다."""
    return display_text(item_text)

def normalize_area_list(value: Any) -> List[str]:
    if value is None:
        return []
    text = normalize_text(value)
    if not text:
        return []
    for sep in [",", ";", "/", "·"]:
        text = text.replace(sep, "|")
    parts = [normalize_area_name(x.strip()) for x in text.split("|") if x.strip()]
    result: List[str] = []
    for p in parts:
        if p not in result:
            result.append(p)
    return result

# -----------------------------------------------------------------------------
# 학교 자료 연결
# -----------------------------------------------------------------------------
def load_school_databases() -> Dict[str, Any]:
    school_info_df, school_info_path, school_info_err = load_optional_csv(CSV_PATHS["school_info"], "학교 정보 자료")
    school_context_df, school_context_path, school_context_err = load_optional_csv(CSV_PATHS["school_context"], "학교 맥락 점수 자료")
    region_context_df, region_context_path, region_context_err = load_optional_csv(CSV_PATHS["region_context"], "지역 맥락 점수 자료")

    errors = [x for x in [school_info_err, school_context_err, region_context_err] if x]

    if school_info_df is None:
        school_info_df = pd.DataFrame(
            [
                {
                    "자치구": "강서구",
                    "교육지원청": "강서양천교육지원청",
                    "학교급": "고등학교",
                    "학교명": "데모고등학교",
                    "학교_주소": "서울특별시 강서구",
                    "학교_학생수": 600,
                    "학교_교원1인당학생수": 12.0,
                    "교육복지우선지원여부_UI": "해당 없음",
                    "위클래스_있음": "있음",
                    "전문상담교사_수": 1,
                    "보건교사_수": 1,
                    "진로상담실_있음": "있음",
                }
            ]
        )
    return {
        "school_info_df": school_info_df,
        "school_context_df": school_context_df,
        "region_context_df": region_context_df,
        "paths": {
            "school_info": school_info_path,
            "school_context": school_context_path,
            "region_context": region_context_path,
        },
        "errors": errors,
    }


def find_matching_row(df: Optional[pd.DataFrame], school_name: str, district: Optional[str] = None) -> Optional[pd.Series]:
    if df is None or df.empty or "학교명" not in df.columns:
        return None
    matched = df[df["학교명"].astype(str) == str(school_name)]
    if district and "자치구" in df.columns:
        matched2 = matched[matched["자치구"].astype(str) == str(district)]
        if not matched2.empty:
            matched = matched2
    if matched.empty:
        return None
    return matched.iloc[0]


def find_region_row(df: Optional[pd.DataFrame], district: str) -> Optional[pd.Series]:
    if df is None or df.empty or "자치구" not in df.columns:
        return None
    matched = df[df["자치구"].astype(str) == str(district)]
    if matched.empty:
        return None
    return matched.iloc[0]


def school_row_to_dict(row: pd.Series) -> Dict[str, Any]:
    def yes(v: Any) -> bool:
        return normalize_text(v) in ["있음", "Y", "예", "1", "True", "true", "해당"]

    def first_number(keys: List[str]) -> Optional[float]:
        for key in keys:
            try:
                if key in row.index:
                    value = get_cell(row, key, None)
                    if value is not None:
                        return value
            except Exception:
                continue
        return None

    # 학교 자료에는 위도/경도 컬럼이 있으나, 기존 코드가 selected_school_info를 만들 때
    # 해당 값을 제외하고 있어 접근성 거리 계산에서 학교 좌표를 찾지 못했다.
    # 여러 가능한 컬럼명을 표준 키와 별칭에 함께 저장해 이후 거리 계산 함수가 안정적으로 읽도록 한다.
    school_lat = first_number(["위도", "학교_위도", "학교위도", "latitude", "lat", "school_latitude", "Y", "y"])
    school_lon = first_number(["경도", "학교_경도", "학교경도", "longitude", "lon", "lng", "school_longitude", "X", "x"])

    return {
        "자치구": normalize_text(row.get("자치구", "")),
        "교육지원청": normalize_text(row.get("교육지원청", "")),
        "학교급": normalize_text(row.get("학교급", "")),
        "학교명": normalize_text(row.get("학교명", "")),
        "학교_주소": normalize_text(row.get("학교_주소", row.get("주소", ""))),
        "대표_전화번호": normalize_text(row.get("대표_전화번호", "")),
        "홈페이지_URL": normalize_text(row.get("홈페이지_URL", "")),
        "학교_학생수": get_cell(row, "학교_학생수", 0),
        "학교_교원1인당학생수": get_cell(row, "학교_교원1인당학생수", 0),
        "교육복지우선지원여부_UI": normalize_text(row.get("교육복지우선지원여부_UI", "")),
        "위클래스_있음": yes(row.get("위클래스_있음", "")),
        "전문상담교사_수": int(get_cell(row, "전문상담교사_수", 0)),
        "보건교사_수": int(get_cell(row, "보건교사_수", 0)),
        "진로상담실_있음": yes(row.get("진로상담실_있음", "")),
        "방과후교과프로그램수": get_cell(row, "방과후교과프로그램수", 0),
        "방과후특기적성프로그램수": get_cell(row, "방과후특기적성프로그램수", 0),
        "위도": school_lat,
        "경도": school_lon,
        "latitude": school_lat,
        "longitude": school_lon,
        "school_latitude": school_lat,
        "school_longitude": school_lon,
    }

# -----------------------------------------------------------------------------
# 학생 데모 데이터
# -----------------------------------------------------------------------------
def classify_action_stage(raw_score: int, red_flag_result: Dict[str, Any]) -> Dict[str, str]:
    if raw_score <= 3:
        score_based_stage = "일상적 관찰"
        score_based_action = "현재 유지"
    elif raw_score <= 7:
        score_based_stage = "주의 및 탐색"
        score_based_action = "교사 면담 및 가벼운 개입"
    else:
        score_based_stage = "심층 파악 필요"
        score_based_action = "2차 상담 질문 활성화 및 전문 상담 검토"

    if red_flag_result.get("urgent_flag"):
        final_action_stage = "심층 파악 필요"
        final_action = "2차 상담 질문 활성화 및 긴급 확인 질문 포함"
    else:
        final_action_stage = score_based_stage
        final_action = score_based_action

    return {
        "score_based_stage": score_based_stage,
        "score_based_action": score_based_action,
        "final_action_stage": final_action_stage,
        "final_action": final_action,
    }


def generate_demo_students(school_name: str) -> pd.DataFrame:
    rows: List[Dict[str, Any]] = []
    patterns = [
        ("일상적 관찰", [0, 0, 0, 0], "특이 신호 없음", False),
        ("주의 및 탐색", [2, 1, 0, 0], "수업 중 무기력, 과제 수행 부족", False),
        ("심층 파악 필요", [1, 4, 1, 0], "또래관계 단절, 우울 표현, 결석 증가", True),
        ("주의 및 탐색", [0, 1, 2, 0], "복장·위생 상태 우려, 돌봄 공백 가능성", False),
        ("심층 파악 필요", [3, 1, 2, 1], "학업 의지 저하, 경제적 어려움, 진로 무기력", False),
        ("주의 및 탐색", [0, 2, 0, 1], "불안 표현, 진로 결정 회피", False),
        ("일상적 관찰", [1, 0, 0, 0], "일시적 과제 미제출", False),
        ("주의 및 탐색", [1, 0, 0, 2], "진로 목표 부재, 활동 참여 저하", False),
    ]
    idx = 1
    for grade in ["1학년", "2학년", "3학년"]:
        for klass in ["1반", "2반", "3반", "4반"]:
            for n in range(1, 9):
                stage, scores, signal, urgent = patterns[(idx + n) % len(patterns)]
                raw_like = sum(scores)
                if urgent:
                    final_stage = "심층 파악 필요"
                elif raw_like >= 5:
                    final_stage = "심층 파악 필요"
                elif raw_like >= 2:
                    final_stage = "주의 및 탐색"
                else:
                    final_stage = "일상적 관찰"
                rows.append(
                    {
                        "학생코드": f"S-{idx:03d}",
                        "이름": f"학생 {idx:03d}",
                        "학교명": school_name,
                        "학년": grade,
                        "반": klass,
                        "학업": scores[0],
                        "심리정서": scores[1],
                        "복지경제": scores[2],
                        "진로": scores[3],
                        "RedFlag": urgent,
                        "주요신호": signal,
                        "최종단계": final_stage,
                        "권장Action": "2차 상담 질문 활성화 및 전문 상담 검토" if final_stage == "심층 파악 필요" else ("교사 면담 및 가벼운 개입" if final_stage == "주의 및 탐색" else "현재 유지"),
                        "담당자": "담임교사" if klass == DEFAULT_CLASS else "학생맞춤통합지원 담당",
                        "기한": str(date.today()) if final_stage != "일상적 관찰" else "-",
                    }
                )
                idx += 1
    return pd.DataFrame(rows)


def get_view_students() -> pd.DataFrame:
    df = st.session_state.students.copy()
    role = st.session_state.role
    if role == ROLE_HOMEROOM:
        df = df[(df["학년"] == st.session_state.homeroom_grade) & (df["반"] == st.session_state.homeroom_class)]
    return df.reset_index(drop=True)

# -----------------------------------------------------------------------------
# 체크리스트 계산 함수
# -----------------------------------------------------------------------------
def get_active_items_df() -> pd.DataFrame:
    items_df = st.session_state.get("checklist_items_df")
    if items_df is None or items_df.empty:
        return pd.DataFrame()
    df = items_df.copy()
    if "active" in df.columns:
        df = df[df["active"].astype(str).str.upper().str.strip() == "Y"]
    for col in ["domain_order", "item_order"]:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce").fillna(999)
    sort_cols = [c for c in ["domain_order", "item_order"] if c in df.columns]
    if sort_cols:
        df = df.sort_values(sort_cols)
    return df.reset_index(drop=True)


def calculate_checklist_scores(items_df: pd.DataFrame, responses: Dict[str, int]) -> Dict[str, Any]:
    active_df = items_df.copy()
    if "active" in active_df.columns:
        active_df = active_df[active_df["active"].astype(str).str.upper().str.strip() == "Y"]

    domain_rows: List[Dict[str, Any]] = []
    for domain in SUPPORT_AREAS:
        sub = active_df[active_df["domain"].map(normalize_area_name) == domain]
        raw = int(sum(int(responses.get(str(row["item_id"]), 0)) for _, row in sub.iterrows()))
        max_score = int(len(sub) * 2)
        scaled = round(raw / max_score * 100, 1) if max_score > 0 else 0.0
        domain_rows.append(
            {
                "지원 영역": domain,
                "domain_raw_score": raw,
                "domain_max_score": max_score,
                "domain_scaled_score": scaled,
            }
        )

    student_raw_score = int(sum(int(v) for k, v in responses.items() if not str(k).startswith("_")))
    student_scaled_score = int(student_raw_score * 5)
    domain_scores = pd.DataFrame(domain_rows)
    primary_areas = get_primary_areas(domain_scores)
    return {
        "student_raw_score": student_raw_score,
        "student_scaled_score": student_scaled_score,
        "domain_scores": domain_scores,
        "primary_areas": primary_areas,
    }


def get_primary_areas(domain_scores: pd.DataFrame) -> List[str]:
    if domain_scores is None or domain_scores.empty:
        return ["현재 뚜렷한 우선 영역 없음"]
    max_score = float(domain_scores["domain_scaled_score"].max())
    if max_score <= 0:
        return ["현재 뚜렷한 우선 영역 없음"]
    return domain_scores.loc[domain_scores["domain_scaled_score"] == max_score, "지원 영역"].tolist()


def detect_red_flags(items_df: pd.DataFrame, responses: Dict[str, int]) -> Dict[str, Any]:
    urgent_items: List[Dict[str, Any]] = []
    for _, row in items_df.iterrows():
        item_id = normalize_text(row.get("item_id"))
        item_code = normalize_text(row.get("item_code"))
        score = int(responses.get(item_id, 0))
        is_red = item_code in ["정서-C", "환경-A"] or item_id in ["EMO_C", "ENV_A"]
        if is_red and score >= 1:
            if item_code == "정서-C" or item_id == "EMO_C":
                mapped_area = "심리정서"
            elif item_code == "환경-A" or item_id == "ENV_A":
                mapped_area = "복지경제"
            else:
                mapped_area = normalize_area_name(row.get("domain", ""))
            urgent_items.append(
                {
                    "item_id": item_id,
                    "item_code": item_code,
                    "item_text": normalize_text(row.get("item_text")),
                    "score": score,
                    "area": mapped_area,
                    "reason": "우선 확인 필요 신호",
                }
            )
    return {
        "urgent_flag": bool(urgent_items),
        "urgent_flag_items": urgent_items,
    }


def parse_rule_ids(value: Any) -> List[str]:
    text = normalize_text(value)
    if not text:
        return []
    for sep in [",", ";", "/"]:
        text = text.replace(sep, "|")
    return [x.strip() for x in text.split("|") if x.strip()]


def row_get_flexible(row: pd.Series, candidates: List[str], default: str = "") -> str:
    for c in candidates:
        if c in row.index:
            value = normalize_text(row.get(c))
            if value:
                return value
    return default


def activate_deep_rules(
    responses: Dict[str, int],
    rule_map_df: Optional[pd.DataFrame],
    deep_rules_df: Optional[pd.DataFrame],
) -> List[Dict[str, Any]]:
    if rule_map_df is None or deep_rules_df is None or rule_map_df.empty or deep_rules_df.empty:
        return []

    if "rule_id" not in deep_rules_df.columns:
        return []

    active_items = {item_id for item_id, score in responses.items() if not str(item_id).startswith("_") and int(score) >= 1}
    active_rule_ids: Dict[str, str] = {}

    # 매핑 파일 기반 활성화
    for _, row in rule_map_df.iterrows():
        item_id = normalize_text(row.get("item_id"))
        item_code = normalize_text(row.get("item_code"))
        related = []
        if "rule_id" in row.index:
            related.extend(parse_rule_ids(row.get("rule_id")))
        if "related_rule_ids" in row.index:
            related.extend(parse_rule_ids(row.get("related_rule_ids")))
        threshold = 1
        if "activation_threshold" in row.index:
            try:
                threshold = int(float(row.get("activation_threshold")))
            except Exception:
                threshold = 1
        item_score = int(responses.get(item_id, 0))
        if item_score >= threshold:
            for rid in related:
                active_rule_ids[rid] = "single"
        if item_code:
            for resp_item_id, resp_score in responses.items():
                if int(resp_score) >= threshold and resp_item_id == item_id:
                    for rid in related:
                        active_rule_ids[rid] = "single"

    # 규칙 파일의 trigger_items 기반 활성화, 복합 규칙 포함
    for _, rule in deep_rules_df.iterrows():
        rid = normalize_text(rule.get("rule_id"))
        trigger_items = parse_rule_ids(rule.get("trigger_items"))
        if not trigger_items:
            continue
        checked_count = sum(1 for item in trigger_items if int(responses.get(item, 0)) >= 1)
        if checked_count == 0:
            continue
        if len(trigger_items) >= 2:
            activation_type = "full_combo" if checked_count == len(trigger_items) else "partial_combo"
        else:
            activation_type = "single"
        previous = active_rule_ids.get(rid)
        if previous == "full_combo":
            continue
        active_rule_ids[rid] = activation_type

    result: List[Dict[str, Any]] = []
    for rid, activation_type in active_rule_ids.items():
        matched = deep_rules_df[deep_rules_df["rule_id"].astype(str) == rid]
        if matched.empty:
            continue
        row = matched.iloc[0]
        rule_title = row_get_flexible(row, ["rule_title", "rule_name"], rid)
        possible_hidden = row_get_flexible(row, ["possible_hidden_factors", "possible_underlying_factors"], "")
        linked_areas = normalize_area_list(row.get("linked_areas"))
        counseling_question_areas = normalize_area_list(row.get("counseling_question_areas"))
        result.append(
            {
                "rule_id": rid,
                "rule_title": rule_title,
                "activation_type": activation_type,
                "surface_signal": row_get_flexible(row, ["surface_signal"], ""),
                "possible_hidden_factors": possible_hidden,
                "deep_guidance_text": row_get_flexible(row, ["deep_guidance_text"], ""),
                "linked_areas": linked_areas,
                "counseling_question_areas": counseling_question_areas,
                "priority": row_get_flexible(row, ["priority"], ""),
                "is_urgent_related": row_get_flexible(row, ["is_urgent_related", "urgent_flag_related"], ""),
            }
        )
    type_order = {"full_combo": 0, "partial_combo": 1, "single": 2}
    result.sort(key=lambda x: (type_order.get(x.get("activation_type", "single"), 9), x.get("rule_id", "")))
    return result


def derive_counseling_consideration_areas(
    domain_scores: pd.DataFrame,
    responses: Dict[str, int],
    red_flag_result: Dict[str, Any],
    active_deep_rules: List[Dict[str, Any]],
    items_df: Optional[pd.DataFrame] = None,
) -> List[Dict[str, Any]]:
    focus: Dict[str, float] = {area: 0.0 for area in SUPPORT_AREAS}
    reasons: Dict[str, List[str]] = {area: [] for area in SUPPORT_AREAS}

    for _, row in domain_scores.iterrows():
        area = normalize_area_name(row["지원 영역"])
        scaled = float(row["domain_scaled_score"])
        if area in focus and scaled > 0:
            focus[area] += scaled * 0.5
            reasons[area].append(f"{area} 영역 1차 체크리스트 환산점수 {scaled:g}점")

    if items_df is not None and not items_df.empty:
        for _, row in items_df.iterrows():
            item_id = normalize_text(row.get("item_id"))
            score = int(responses.get(item_id, 0))
            area = normalize_area_name(row.get("domain"))
            if area in focus and score == 2:
                focus[area] += 10
                reasons[area].append("해당 영역의 관찰 신호가 뚜렷하게 확인됨")
            if area in focus and score >= 1:
                text = clean_checklist_item_label(row.get("item_text"))
                short = text[:32] + ("..." if len(text) > 32 else "")
                reasons[area].append(f"관찰 신호: {short}")

    for item in red_flag_result.get("urgent_flag_items", []):
        area = normalize_area_name(item.get("area"))
        if area in focus:
            focus[area] += 20
            reasons[area].append("우선 확인 필요 신호")

    for rule in active_deep_rules:
        title = display_text(rule.get("rule_title", rule.get("rule_id", "심층 유도 분석")))
        for area in rule.get("linked_areas", []):
            area = normalize_area_name(area)
            if area in focus:
                focus[area] += 15
                reasons[area].append(f"심층 유도 분석 연결 영역: {title}")
        for area in rule.get("counseling_question_areas", []):
            area = normalize_area_name(area)
            if area in focus:
                focus[area] += 15
                reasons[area].append(f"상담 질문 생성 시 함께 확인 필요: {title}")

    result: List[Dict[str, Any]] = []
    for area, score in focus.items():
        if score <= 0:
            continue
        if score >= 60:
            level = "필수 확인"
        elif score >= 35:
            level = "함께 확인"
        else:
            level = "보조 확인"
        unique_reasons = []
        for r in reasons[area]:
            if r and r not in unique_reasons:
                unique_reasons.append(r)
        result.append(
            {
                "area": area,
                "focus_score": round(score, 1),
                "priority_level": level,
                "reasons": unique_reasons[:6],
            }
        )

    if red_flag_result.get("urgent_flag"):
        result.append(
            {
                "area": "긴급확인",
                "focus_score": 100,
                "priority_level": "필수 확인",
                "reasons": ["우선 확인 필요 신호가 포함되어 안전·정서·환경 확인 질문을 우선 포함"],
            }
        )

    level_order = {"필수 확인": 0, "함께 확인": 1, "보조 확인": 2}
    result.sort(key=lambda x: (level_order.get(x["priority_level"], 9), -float(x["focus_score"])))
    return result


def calculate_school_bonus(score: float, apply_context: bool) -> float:
    if not apply_context:
        return 0.0
    if score >= 90:
        return 10.0
    if score >= 75:
        return 7.0
    if score >= 50:
        return 3.0
    return 0.0


def calculate_region_bonus(score: float, apply_context: bool) -> float:
    if not apply_context:
        return 0.0
    if score >= 90:
        return 5.0
    if score >= 75:
        return 3.5
    if score >= 50:
        return 1.5
    return 0.0



def score_based_stage_from_raw(raw_score: int) -> Tuple[str, str]:
    if raw_score <= 3:
        return "일상적 관찰", "현재 유지"
    if raw_score <= 7:
        return "주의 및 탐색", "교사 면담 및 가벼운 개입"
    return "심층 파악 필요", "2차 상담 질문 활성화 및 전문 상담 검토"


def context_adjustment_not_applied_reason(raw_score: int, red_flag: bool) -> str:
    if red_flag:
        return "우선 확인 필요 신호 우선 적용"
    if raw_score <= 4:
        return "원점수 0~4점 구간"
    if raw_score >= 8:
        return "원점수 8점 이상으로 이미 심층 파악 기준 도달"
    return "-"


def determine_final_action_stage(raw_score: int, scaled_score: float, red_flag: bool, context_check_score: float, context_adjustment_applied: bool) -> Dict[str, Any]:
    score_stage, score_action = score_based_stage_from_raw(raw_score)
    if red_flag:
        return {
            "score_based_stage": score_stage,
            "score_based_action": score_action,
            "final_action_stage": "심층 파악 필요",
            "final_action": "2차 상담 질문 활성화 및 우선 확인 질문 포함",
            "final_action_reason": "우선 확인 필요 신호가 포함되어 있습니다.",
            "activate_counseling_form": True,
        }
    if raw_score >= 8:
        return {
            "score_based_stage": score_stage,
            "score_based_action": score_action,
            "final_action_stage": "심층 파악 필요",
            "final_action": "2차 상담 질문 활성화 및 전문 상담 검토",
            "final_action_reason": "1차 체크리스트 원점수가 심층 파악 기준에 도달했습니다.",
            "activate_counseling_form": True,
        }
    if 5 <= raw_score <= 7:
        if context_adjustment_applied and context_check_score >= 40:
            return {
                "score_based_stage": score_stage,
                "score_based_action": score_action,
                "final_action_stage": "심층 파악 권고",
                "final_action": "2차 상담 질문 생성 권고",
                "final_action_reason": "체크리스트 기준으로는 주의 및 탐색 단계이나, 학교·지역 지원여건을 고려할 때 2차 상담 질문으로 한 번 더 확인하는 것이 권장됩니다.",
                "activate_counseling_form": True,
            }
        return {
            "score_based_stage": score_stage,
            "score_based_action": score_action,
            "final_action_stage": "주의 및 탐색",
            "final_action": "담임교사 면담과 추가 관찰 우선",
            "final_action_reason": "일부 지원 신호가 관찰되지만, 현재는 담임교사의 면담과 추가 관찰을 우선 권장합니다.",
            "activate_counseling_form": False,
        }
    if raw_score == 4:
        return {
            "score_based_stage": score_stage,
            "score_based_action": score_action,
            "final_action_stage": "주의 및 탐색",
            "final_action": "교사 면담 및 가벼운 개입",
            "final_action_reason": "일부 지원 신호가 관찰됩니다. 단, 맥락 보정 적용 구간에는 해당하지 않아 담임교사의 면담과 추가 관찰을 우선 권장합니다.",
            "activate_counseling_form": False,
        }
    return {
        "score_based_stage": score_stage,
        "score_based_action": score_action,
        "final_action_stage": "일상적 관찰",
        "final_action": "현재 유지",
        "final_action_reason": "현재 1차 체크리스트 기준으로는 뚜렷한 지원 신호가 높지 않습니다.",
        "activate_counseling_form": False,
    }


def calculate_context_result(
    primary_areas: List[str],
    student_raw_score: int,
    student_scaled_score: float,
    red_flag_result: Dict[str, Any],
    selected_school_context: Optional[pd.Series],
    selected_region_context: Optional[pd.Series],
) -> Dict[str, Any]:
    usable_primary = [a for a in primary_areas if a in SUPPORT_AREAS]
    applied_area = usable_primary[0] if usable_primary else "-"
    school_col = AREA_TO_SCHOOL_CONTEXT_COL.get(applied_area, "")
    region_col = AREA_TO_REGION_CONTEXT_COL.get(applied_area, "")
    school_context_score = get_cell(selected_school_context, school_col, 0) if school_col else 0.0
    region_context_score = get_cell(selected_region_context, region_col, 0) if region_col else 0.0

    red_flag = bool(red_flag_result.get("urgent_flag"))
    context_adjustment_applied = (5 <= int(student_raw_score) <= 7) and not red_flag
    context_adjustment_reason = "원점수 5~7점 구간이며 우선 확인 필요 신호가 없어 맥락 보정을 적용합니다." if context_adjustment_applied else context_adjustment_not_applied_reason(int(student_raw_score), red_flag)

    school_candidate = calculate_school_bonus(school_context_score, True)
    region_candidate = calculate_region_bonus(region_context_score, True)
    school_applied = school_candidate if context_adjustment_applied else 0.0
    region_applied = region_candidate if context_adjustment_applied else 0.0
    context_check_score = round(float(student_scaled_score) + school_applied + region_applied, 1)
    stage = determine_final_action_stage(int(student_raw_score), float(student_scaled_score), red_flag, context_check_score, context_adjustment_applied)

    reason = "-" if context_adjustment_applied else context_adjustment_reason
    context_rows = [
        {
            "구분": "학교 맥락",
            "적용 영역": applied_area,
            "원 맥락 점수": round(school_context_score, 1),
            "산출 가능한 보정점수": school_candidate,
            "실제 적용 보정점수": school_applied,
            "적용 여부": "적용" if context_adjustment_applied else "미적용",
            "미적용 사유": reason,
        },
        {
            "구분": "지역 맥락",
            "적용 영역": applied_area,
            "원 맥락 점수": round(region_context_score, 1),
            "산출 가능한 보정점수": region_candidate,
            "실제 적용 보정점수": region_applied,
            "적용 여부": "적용" if context_adjustment_applied else "미적용",
            "미적용 사유": reason,
        },
    ]
    return {
        "applied_area": applied_area,
        "context_adjustment_applied": context_adjustment_applied,
        "context_adjustment_reason": context_adjustment_reason,
        "school_context_score": round(school_context_score, 1),
        "region_context_score": round(region_context_score, 1),
        "school_context_bonus_candidate": school_candidate,
        "region_context_bonus_candidate": region_candidate,
        "school_context_bonus_applied": school_applied,
        "region_context_bonus_applied": region_applied,
        "context_check_score": context_check_score,
        "context_table": pd.DataFrame(context_rows),
        "context_missing": selected_school_context is None or selected_region_context is None,
        **stage,
    }


def build_counseling_payload(
    first_check_result: Dict[str, Any],
    red_flag_result: Dict[str, Any],
    context_result: Dict[str, Any],
    active_deep_rules: List[Dict[str, Any]],
    counseling_consideration_areas: List[Dict[str, Any]],
    stage_result: Dict[str, Any],
) -> Dict[str, Any]:
    domain_scores_payload = {}
    for _, row in first_check_result["domain_scores"].iterrows():
        area = row["지원 영역"]
        domain_scores_payload[area] = {
            "domain_raw_score": int(row["domain_raw_score"]),
            "domain_max_score": int(row["domain_max_score"]),
            "domain_scaled_score": float(row["domain_scaled_score"]),
        }

    payload = {
        "first_check_result": {
            "raw_score": int(first_check_result["student_raw_score"]),
            "scaled_score": int(first_check_result["student_scaled_score"]),
            "score_based_stage": stage_result["score_based_stage"],
            "final_action_stage": stage_result["final_action_stage"],
            "final_action_reason": stage_result.get("final_action_reason", ""),
            "activate_counseling_form": bool(stage_result.get("activate_counseling_form", False)),
            "domain_scores": domain_scores_payload,
            "primary_areas": first_check_result["primary_areas"],
        },
        "red_flag_result": {
            "urgent_flag": bool(red_flag_result.get("urgent_flag", False)),
            "urgent_flag_items": red_flag_result.get("urgent_flag_items", []),
        },
        "context_result": {
            "context_adjustment_applied": bool(context_result.get("context_adjustment_applied", False)),
            "context_adjustment_reason": context_result.get("context_adjustment_reason"),
            "school_context_score": context_result.get("school_context_score"),
            "region_context_score": context_result.get("region_context_score"),
            "school_context_bonus_candidate": context_result.get("school_context_bonus_candidate"),
            "region_context_bonus_candidate": context_result.get("region_context_bonus_candidate"),
            "school_context_bonus_applied": context_result.get("school_context_bonus_applied"),
            "region_context_bonus_applied": context_result.get("region_context_bonus_applied"),
            "context_check_score": context_result.get("context_check_score"),
        },
        "active_deep_rules": [
            {
                "rule_id": r.get("rule_id"),
                "rule_title": r.get("rule_title"),
                "activation_type": r.get("activation_type"),
                "surface_signal": r.get("surface_signal"),
                "possible_hidden_factors": r.get("possible_hidden_factors"),
                "deep_guidance_text": r.get("deep_guidance_text"),
                "linked_areas": r.get("linked_areas"),
                "counseling_question_areas": r.get("counseling_question_areas"),
            }
            for r in active_deep_rules
        ],
        "counseling_consideration_areas": counseling_consideration_areas,
        "instruction_for_next_step": "이 결과를 기반으로 교육청 체크리스트를 참고하여 2차 상담 질문을 생성한다.",
    }
    return payload

# -----------------------------------------------------------------------------
# 렌더링 함수
# -----------------------------------------------------------------------------
def render_header() -> None:
    school = st.session_state.selected_school_info
    role = st.session_state.role
    scope_text = f"{st.session_state.homeroom_grade} {st.session_state.homeroom_class}만 조회" if role == ROLE_HOMEROOM else "전교 학생 조회"
    st.markdown(
        f"""
        <div class="edutop">
            <div class="edu-logo">
                <span class="edu-symbol">학</span>
                <div>
                    <div>학생맞춤통합지원 AI</div>
                    <div style="font-size:.74rem;color:#64748b;font-weight:700;">교사용 지원 신호 포털</div>
                </div>
            </div>
            <div class="edu-main">
                <div class="edu-bluebar">
                    <div>
                        <span class="edu-school-badge">{school.get('학교명', '학교 미선택')}</span>
                        {school.get('교육지원청', '')} · {school.get('자치구', '')}
                    </div>
                    <div class="edu-actions">
                        <span>로그아웃</span><span>튜토리얼</span><span>사용자지원</span>
                    </div>
                </div>
                <div class="rolebar">
                    <div>현재 역할: <span class="role-pill">{role}</span></div>
                    <div class="small-muted">조회 범위: {scope_text}</div>
                </div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_role_switch() -> None:
    role = st.radio(
        "역할 전환",
        [ROLE_HOMEROOM, ROLE_COORDINATOR],
        index=[ROLE_HOMEROOM, ROLE_COORDINATOR].index(st.session_state.role),
        horizontal=True,
        help="담임교사는 본인 반 학생만, 학생맞춤통합지원담당교원은 전교 학생을 볼 수 있습니다.",
    )
    st.session_state.role = role


def render_page_title(title: str, subtitle: str) -> None:
    st.markdown(f"<div class='page-title'>{title}</div>", unsafe_allow_html=True)
    st.markdown(f"<div class='page-subtitle'>{subtitle}</div>", unsafe_allow_html=True)


def render_workflow_section_header(number: str, title: str, subtitle: str = "") -> None:
    st.markdown(
        f"""
        <div class="workflow-section-header">
            <div class="workflow-section-title"><span class="workflow-section-kicker">{html.escape(str(number))}</span>{html.escape(str(title))}</div>
            <div class="workflow-section-subtitle">{html.escape(str(subtitle))}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def area_chips_html(areas: List[str], css_class: str = "area-chip-main") -> str:
    clean = display_area_list([a for a in areas if normalize_area_name(a) != "긴급확인"])
    if not clean:
        clean = ["현재 뚜렷한 우선 영역 없음"]
    chips = "".join(f"<span class='{css_class}'>{html.escape(display_text(a))}</span>" for a in dict.fromkeys(clean))
    return f"<div class='area-chip-row'>{chips}</div>"


def area_help_text(prefix: str, areas: List[str], empty_text: str = "추가로 표시할 영역 없음") -> str:
    clean = display_area_list([a for a in areas if normalize_area_name(a) != "긴급확인"])
    if clean:
        return f"{display_text(prefix)}: {', '.join(clean)}"
    return display_text(empty_text)


def get_checklist_direct_and_related_areas(first_check_result: Dict[str, Any], counseling_areas: List[Dict[str, Any]], active_deep_rules: Optional[List[Dict[str, Any]]] = None) -> Tuple[List[str], List[str]]:
    """체크리스트에서 직접 점수가 나온 영역과 심층분석으로 함께 고려할 영역을 분리한다."""
    direct: List[str] = []
    domain_scores = first_check_result.get("domain_scores", {})
    if isinstance(domain_scores, dict):
        for area, score in domain_scores.items():
            try:
                val = float(score.get("domain_scaled_score", score) if isinstance(score, dict) else score)
            except Exception:
                val = 0.0
            if area in SUPPORT_AREAS and val > 0:
                direct.append(area)
    if not direct:
        for area in first_check_result.get("primary_areas", []) or []:
            if area in SUPPORT_AREAS and area not in direct:
                direct.append(area)

    related: List[str] = []
    for item in counseling_areas or []:
        area = normalize_area_name(item.get("area"))
        if area in SUPPORT_AREAS and area not in direct and area not in related:
            related.append(area)
    for rule in active_deep_rules or []:
        for area in normalize_target_areas(rule.get("linked_areas", [])) + normalize_target_areas(rule.get("counseling_question_areas", [])):
            if area in SUPPORT_AREAS and area not in direct and area not in related:
                related.append(area)
    return list(dict.fromkeys(direct)), list(dict.fromkeys(related))


def metric_card(label: str, value: str, help_text: str = "") -> str:
    label_t = display_text(label)
    value_t = display_text(value)
    help_t = display_text(help_text)
    return f"""
    <div class="metric-card">
        <div class="metric-label">{html.escape(label_t)}</div>
        <div class="metric-value">{html.escape(value_t)}</div>
        <div class="metric-help">{html.escape(help_t)}</div>
    </div>
    """


def analysis_summary_card(status: str, summary_text: str) -> str:
    return f"""
    <div class="analysis-summary-card">
        <div class="analysis-summary-label">지원 검토 안내</div>
        <div class="analysis-summary-status">{html.escape(display_text(status or '-'))}</div>
        <div class="analysis-summary-text">{html.escape(display_text(summary_text or '상담 결과를 바탕으로 추가 검토가 필요합니다.'))}</div>
    </div>
    """


def resource_detail_table_html(rows: List[Tuple[str, Any]]) -> str:
    body = []
    for label, value in rows:
        val = "-" if value is None or str(value).strip() == "" else display_text(value)
        body.append(
            f"<tr><th>{html.escape(display_text(label))}</th><td>{html.escape(val)}</td></tr>"
        )
    return "<table class='resource-detail-table'>" + "".join(body) + "</table>"




def user_friendly_generation_error(error: Any, task_label: str) -> str:
    """내부 검증 메시지를 실제 사용자에게 부드러운 안내문으로 바꾼다."""
    err = str(error or "").strip()
    if not err:
        return f"{task_label} 처리 중 일시적인 오류가 발생했습니다. 잠시 후 다시 시도해 주세요."
    api_markers = [
        "Gemini API 사용량", "무료 등급", "사용량 제한", "API 키", "접근 권한", "프로젝트", "모델 서버", "일시적으로 혼잡", "GEMINI_API_KEY",
    ]
    if any(m in err for m in api_markers):
        return err
    validation_markers = [
        "금지 표현", "검증", "JSON", "필수", "허용 범위", "질문 개수", "recommended_questions", "support_needed", "primary_area", "decision_items",
    ]
    if any(m in err for m in validation_markers):
        return (
            f"{task_label} 결과를 정리하는 과정에서 일부 표현을 다시 다듬어야 합니다. "
            "잠시 후 다시 시도하거나 입력 내용을 조금 더 구체적으로 작성해 주세요."
        )
    return f"{task_label} 처리 중 일시적인 오류가 발생했습니다. 잠시 후 다시 시도해 주세요."


def stage_badge(stage: str) -> str:
    if stage in ["심층 파악 필요", "심층 파악 권고"]:
        cls = "badge-deep"
    elif stage == "주의 및 탐색":
        cls = "badge-watch"
    else:
        cls = "badge-normal"
    return f"<span class='badge {cls}'>{stage}</span>"


def dots(score: int, max_score: int = 4) -> str:
    if score <= 0:
        return "<span class='dots dots-zero'>○○○○</span>"
    filled = "●" * min(score, max_score)
    empty = "○" * max(0, max_score - score)
    return f"<span class='dots'>{filled}<span class='dots-zero'>{empty}</span></span>"


def domain_grid(row: pd.Series) -> str:
    cells = "".join(
        f"""
        <div class="mini-domain">
            <div class="domain-name">{html.escape(display_area_name(area))}</div>
            {dots(int(row.get(area, 0)))}
        </div>
        """
        for area in SUPPORT_AREAS
    )
    return f"<div class='mini-grid'>{cells}</div>"


def render_student_card(row: pd.Series) -> None:
    stage = normalize_text(row.get("최종단계", "일상적 관찰"))
    cls = "risk-deep" if stage in ["심층 파악 필요", "심층 파악 권고"] else ("risk-watch" if stage == "주의 및 탐색" else "risk-normal")
    top_area = max(SUPPORT_AREAS, key=lambda a: int(row.get(a, 0)))
    if sum(int(row.get(a, 0)) for a in SUPPORT_AREAS) == 0:
        top_area = "-"
    st.markdown(
        f"""
        <div class="student-card {cls}">
            <div style="display:flex;justify-content:space-between;align-items:center;gap:10px;">
                <div style="font-weight:900;font-size:1.02rem;color:#0f172a;">
                    {stage_badge(stage)} {row.get('이름')} <span class="small-muted">({row.get('학생코드')})</span>
                </div>
                <div class="small-muted">{row.get('학년')} {row.get('반')} · 우선 영역: <span style="font-weight:900;">{display_area_name(top_area) if top_area != '-' else '-'}</span></div>
            </div>
            {domain_grid(row)}
            <div style="margin-top:10px;color:#334155;font-size:.9rem;">
                <span style="font-weight:900;">주요 신호</span>: {display_text(row.get('주요신호'))}<br>
                <span style="font-weight:900;">권장 Action</span>: {display_text(row.get('권장Action'))}
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_school_info_cards() -> None:
    school = st.session_state.selected_school_info
    cols = st.columns(5)
    values = [
        ("학교", school.get("학교명", "-"), school.get("학교급", "")),
        ("학생 수", f"{int(school.get('학교_학생수', 0)):,}명" if school.get("학교_학생수", 0) else "-", "학교 자료 기준"),
        ("교원 1인당 학생수", f"{school.get('학교_교원1인당학생수', 0):.1f}명" if school.get("학교_교원1인당학생수", 0) else "-", "학교 여건"),
        ("위클래스", "있음" if school.get("위클래스_있음") else "없음", f"상담교사 {school.get('전문상담교사_수', 0)}명"),
        ("진로상담실", "있음" if school.get("진로상담실_있음") else "없음", f"보건교사 {school.get('보건교사_수', 0)}명"),
    ]
    for col, (label, value, help_text) in zip(cols, values):
        with col:
            st.markdown(metric_card(label, str(value), help_text), unsafe_allow_html=True)

def render_summary_cards(result: Dict[str, Any]) -> None:
    stage = result.get("final_action_stage", "-")
    action = result.get("final_action", "-")
    areas = result.get("support_areas", []) or []
    urgent = "있음" if result.get("urgent_flag") else "없음"
    cards = [
        ("최종 안내 단계", stage, result.get("final_action_reason", "")),
        ("권장 조치", action, "교사 검토용 안내"),
        ("복합 지원 영역", ", ".join(areas) if areas else "현재 뚜렷한 우선 영역 없음", "상담과 지원 연계 시 함께 볼 영역"),
        ("우선 확인 필요 신호", urgent, "해당 신호가 있으면 추가 확인을 우선 고려"),
    ]
    cols = st.columns(4)
    for col, (label, value, help_text) in zip(cols, cards):
        with col:
            st.markdown(metric_card(label, value, help_text), unsafe_allow_html=True)


def render_compact_checklist_result(context_result: Dict[str, Any], red_flag_result: Dict[str, Any], counseling_areas: List[Dict[str, Any]]) -> None:
    final_stage = context_result.get("final_action_stage", "-")
    support_areas = [x.get("area") for x in counseling_areas if x.get("area") and x.get("area") != "긴급확인"]
    # 중복 제거
    seen = []
    for a in support_areas:
        if a not in seen:
            seen.append(a)
    summary_for_cards = {
        "final_action_stage": final_stage,
        "final_action_reason": context_result.get("final_action_reason", ""),
        "final_action": context_result.get("final_action", "-"),
        "urgent_flag": red_flag_result.get("urgent_flag", False),
        "support_areas": seen,
    }
    render_summary_cards(summary_for_cards)
    if final_stage == "일상적 관찰":
        st.info("현재는 평소 관찰과 라포 형성을 유지하는 단계입니다.")
    elif final_stage == "주의 및 탐색":
        st.warning("일부 지원 신호가 관찰됩니다. 담임교사의 가벼운 면담과 추가 관찰을 권장합니다.")
    elif final_stage == "심층 파악 권고":
        st.warning("2차 상담 질문을 통해 학생의 최근 변화와 어려움을 한 번 더 확인하는 것이 권장됩니다.")
    else:
        st.error("심층 파악이 필요합니다. 다음 단계에서 2차 상담 질문을 생성할 수 있습니다.")


def render_domain_score_table(domain_scores: pd.DataFrame, primary_areas: Optional[List[str]] = None) -> None:
    if domain_scores is None or domain_scores.empty:
        st.info("영역별 점수를 표시할 수 없습니다.")
        return
    primary_areas = primary_areas or []
    table = domain_scores.copy()
    table["우선 영역 여부"] = table["지원 영역"].apply(lambda x: "우선" if x in primary_areas else "-")
    table["지원 영역"] = table["지원 영역"].apply(display_area_name)
    table = table.rename(
        columns={
            "domain_raw_score": "원점수",
            "domain_max_score": "최대점수",
            "domain_scaled_score": "환산점수",
        }
    )[["지원 영역", "원점수", "최대점수", "환산점수", "우선 영역 여부"]]
    st.dataframe(table, use_container_width=True, hide_index=True)


def render_red_flag_section(red_flag_result: Dict[str, Any]) -> None:
    st.markdown("<div class='panel'><div class='panel-title'>우선 확인 필요 신호 결과</div>", unsafe_allow_html=True)
    if not red_flag_result.get("urgent_flag"):
        st.info("우선 확인 필요 신호가 확인되지 않았습니다.")
    else:
        st.markdown(
            "<div class='warning-callout'>긴급 확인 관련 신호가 포함되어 2차 상담 질문 생성 시 안전·정서·환경 확인 질문을 우선 포함합니다.</div>",
            unsafe_allow_html=True,
        )
        table = pd.DataFrame(red_flag_result.get("urgent_flag_items", []))
        if not table.empty:
            show_cols = [c for c in ["item_code", "item_text", "score", "area", "reason"] if c in table.columns]
            st.dataframe(table[show_cols], use_container_width=True, hide_index=True)
    st.markdown("</div>", unsafe_allow_html=True)


def render_deep_rule_cards(active_deep_rules: List[Dict[str, Any]]) -> None:
    st.markdown("<div class='panel'><div class='panel-title'>심층 유도 분석 결과</div>", unsafe_allow_html=True)
    if not active_deep_rules:
        st.info("활성화된 심층 유도 분석이 없습니다.")
    for i, rule in enumerate(active_deep_rules, start=1):
        linked = ", ".join(display_area_list(rule.get("linked_areas", []))) or "-"
        st.markdown(
            f"""
            <div class="recommend-card">
                <div style="font-weight:900;"><span class="recommend-rank">{i}</span>{display_text(rule.get('rule_title', rule.get('rule_id')))}</div>
                <div class="small-muted">활성화 유형: {display_text(rule.get('activation_type'))}</div>
                <table class="info-table" style="margin-top:8px;">
                    <tr><th>표면 신호</th><td>{display_text(rule.get('surface_signal', '-'))}</td></tr>
                    <tr><th>가능한 이면 변인</th><td>{display_text(rule.get('possible_hidden_factors', '-'))}</td></tr>
                    <tr><th>심층 유도 서술</th><td>{display_text(rule.get('deep_guidance_text', '-'))}</td></tr>
                    <tr><th>연결 지원 영역</th><td>{linked}</td></tr>
                </table>
            </div>
            """,
            unsafe_allow_html=True,
        )
    st.markdown("</div>", unsafe_allow_html=True)


def render_counseling_area_table(counseling_areas: List[Dict[str, Any]]) -> None:
    st.markdown("<div class='panel'><div class='panel-title'>상담지 생성 시 고려 영역</div>", unsafe_allow_html=True)
    if not counseling_areas:
        st.info("상담지 생성 시 추가로 고려할 영역이 없습니다.")
    else:
        table = pd.DataFrame(
            [
                {
                    "고려 영역": x["area"],
                    "확인 우선도": x["priority_level"],
                    "focus_score": x["focus_score"],
                    "포함 근거": " / ".join(x.get("reasons", [])),
                }
                for x in counseling_areas
            ]
        )
        st.dataframe(table, use_container_width=True, hide_index=True)
    st.markdown("</div>", unsafe_allow_html=True)



def render_context_table(context_result: Dict[str, Any]) -> None:
    st.markdown("<div class='panel'><div class='panel-title'>맥락 반영 심층확인 점수 표</div>", unsafe_allow_html=True)
    if context_result.get("context_missing"):
        st.warning("학교 또는 지역 맥락 점수를 찾을 수 없어 참고 보정은 0점으로 처리했습니다.")
    table = context_result.get("context_table")
    if isinstance(table, pd.DataFrame) and not table.empty:
        st.dataframe(table, use_container_width=True, hide_index=True)
    if context_result.get("context_adjustment_applied"):
        st.caption("체크리스트 원점수 5~7점 구간에서만 학교·지역 맥락을 심층확인 보조 지표로 반영합니다.")
    else:
        st.caption("학교·지역 맥락 점수는 학생 개인의 지원 필요도를 직접 높이는 점수가 아니며, 이번 결과에서는 행동 단계 판단에 적용하지 않았습니다.")
    st.markdown("</div>", unsafe_allow_html=True)


# -----------------------------------------------------------------------------
# v3: Gemini, 상담 질문, 상담 메모 구조화, 기관 추천, 문서 생성
# -----------------------------------------------------------------------------

def load_json_with_fallback(paths: Iterable[Path]) -> Dict[str, Any]:
    path = first_existing_path(paths)
    if path is None:
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8-sig"))
    except Exception:
        try:
            return json.loads(path.read_text(encoding="cp949"))
        except Exception:
            return {}


def load_official_checklist_reference() -> Tuple[Optional[pd.DataFrame], Optional[str]]:
    path = first_existing_path(CSV_PATHS.get("official_checklist", []))
    if path is None:
        return None, "official_counseling_checklist_reference_v1.csv 파일을 찾을 수 없습니다."
    try:
        return load_csv_with_fallback(path), None
    except Exception as exc:
        return None, f"공식 체크리스트 참고 CSV를 읽지 못했습니다: {exc}"


def get_official_checklist_context(
    official_df: pd.DataFrame,
    counseling_consideration_areas: List[Dict[str, Any]],
    red_flag_result: Dict[str, Any],
    max_items: int = 12,
) -> List[Dict[str, Any]]:
    if official_df is None or official_df.empty:
        return []
    df = official_df.copy()
    for c in ["mapped_area_primary", "mapped_area_secondary", "urgent_check_reference", "question_focus", "suggested_question_angle"]:
        if c not in df.columns:
            df[c] = ""
    areas = [normalize_area_name(x.get("area", "")) for x in counseling_consideration_areas if x.get("area")]
    areas = [a for a in areas if a]
    required_areas = [normalize_area_name(x.get("area", "")) for x in counseling_consideration_areas if x.get("priority_level") == "필수 확인"]

    def row_match(row: pd.Series) -> bool:
        primary = normalize_area_name(row.get("mapped_area_primary", ""))
        secondary = normalize_area_list(row.get("mapped_area_secondary", ""))
        if primary in areas or any(a in areas for a in secondary):
            return True
        if primary == "공통":
            return True
        if red_flag_result.get("urgent_flag") and str(row.get("urgent_check_reference", "")).upper().strip() == "Y":
            return True
        return False

    filtered = df[df.apply(row_match, axis=1)].copy()
    if filtered.empty:
        filtered = df[df.get("mapped_area_primary", "").astype(str).map(normalize_area_name).isin(["공통"] + SUPPORT_AREAS)].head(5).copy()

    def priority_score(row: pd.Series) -> int:
        score = 0
        primary = normalize_area_name(row.get("mapped_area_primary", ""))
        secondary = normalize_area_list(row.get("mapped_area_secondary", ""))
        if primary in required_areas or any(a in required_areas for a in secondary):
            score += 50
        if str(row.get("urgent_check_reference", "")).upper().strip() == "Y":
            score += 30
        if normalize_text(row.get("question_focus")):
            score += 5
        if normalize_text(row.get("suggested_question_angle")):
            score += 5
        return score

    filtered["_priority"] = filtered.apply(priority_score, axis=1)
    sort_cols = ["_priority"]
    ascending = [False]
    for col in ["section_order", "item_order"]:
        if col in filtered.columns:
            sort_cols.append(col)
            ascending.append(True)
    filtered = filtered.sort_values(sort_cols, ascending=ascending).head(max_items)

    out: List[Dict[str, Any]] = []
    for _, r in filtered.iterrows():
        out.append(
            {
                "ref_id": normalize_text(r.get("ref_id")),
                "criterion_type": normalize_text(r.get("criterion_type")),
                "mapped_area_primary": normalize_text(r.get("mapped_area_primary")),
                "mapped_area_secondary": normalize_text(r.get("mapped_area_secondary")),
                "urgent_check_reference": normalize_text(r.get("urgent_check_reference")),
                "official_item_text": normalize_text(r.get("official_item_text")),
                "question_focus": normalize_text(r.get("question_focus")),
                "suggested_question_angle": normalize_text(r.get("suggested_question_angle")),
                "teacher_caution": normalize_text(r.get("teacher_caution")),
            }
        )
    return out


def _safe_json(obj: Any) -> str:
    return json.dumps(obj, ensure_ascii=False, indent=2, default=str)


# ------------------------------ 2차 상담 질문 생성 ------------------------------
def build_counseling_question_system_prompt() -> str:
    return """
너는 학생맞춤통합지원 업무를 돕는 교사용 AI 상담 질문 생성 보조도구이다.
너의 역할은 1차 체크리스트 결과, 심층 유도 분석, 공식 교육청 체크리스트 참고 항목을 바탕으로 교사가 학생과 2차 상담을 진행할 때 사용할 수 있는 상담 질문을 추천하는 것이다.
학생을 진단하거나 판정하지 말고, 교사의 최종 판단을 대체하지 않는다.
낙인 표현, 의학적·법적 진단, 취조형 질문, 비교·죄책감 유발, 위협 표현을 사용하지 않는다.
공식 체크리스트 문항을 그대로 복사하지 말고 학생에게 자연스럽게 물어볼 수 있는 개방형 상담 질문으로 바꾼다.
출력은 반드시 지정된 JSON 형식으로만 작성한다.
""".strip()


def build_counseling_question_user_prompt(payload: Dict[str, Any]) -> str:
    return f"""
아래 자료를 바탕으로 학생맞춤통합지원 2차 상담 질문을 생성하라.
이번 상담지는 점수형 체크리스트가 아니라 교사가 학생과 대화할 때 활용할 수 있는 상담 질문 추천지이다.

[1차 체크리스트 결과]
{_safe_json(payload.get('first_check_result'))}

[Red Flag 또는 우선확인 신호 결과]
{_safe_json(payload.get('red_flag_result'))}

[활성화된 심층 유도 분석]
{_safe_json(payload.get('active_deep_rules'))}

[상담 질문 생성 시 고려할 영역]
{_safe_json(payload.get('counseling_consideration_areas'))}

[공식 교육청 체크리스트 참고 항목]
{_safe_json(payload.get('official_checklist_context'))}

중요 지침:
1. 교사 관찰 메모는 이번 입력에 포함되지 않는다.
2. 입력에 없는 학생의 구체적 사정, 가정환경, 질병, 피해 경험, 심리 상태를 만들어내지 않는다.
3. 필수 확인 영역은 반드시 질문에 포함한다.
4. 우선 확인 필요 신호가 있으면 linked_area가 "긴급확인"인 질문을 최소 1개 포함한다.
5. 질문은 총 5~8개 생성한다.
6. 각 질문마다 질문 목적, 연결 영역, 근거, 교사 유의사항을 함께 작성한다.
7. 출력은 JSON만 작성한다.

출력 형식:
{{
  "counseling_focus_summary": "이번 상담에서 중점적으로 확인할 내용 요약",
  "question_generation_basis": {{
    "primary_areas": ["학업"],
    "secondary_areas": ["진로", "심리정서"],
    "urgent_check_needed": false,
    "basis_summary": "질문 생성 근거 요약"
  }},
  "recommended_questions": [
    {{
      "question_id": "Q1",
      "question": "상담 질문 문장",
      "purpose": "이 질문을 하는 목적",
      "linked_area": "진로 / 학업 / 심리정서 / 복지경제 / 긴급확인 중 하나",
      "based_on": ["1차 체크리스트 문항", "심층 유도 분석 rule_id", "공식 체크리스트 참고 항목"],
      "teacher_caution": "교사가 질문할 때 주의할 점",
      "follow_up_if_needed": "학생 답변에 따라 추가로 확인할 수 있는 내용"
    }}
  ],
  "areas_to_confirm": [
    {{"area": "심리정서", "priority": "필수 확인 / 함께 확인 / 보조 확인", "reason": "이유"}}
  ],
  "urgent_check_guidance": {{"urgent_check_needed": false, "guidance": ""}},
  "teacher_recording_guide": ["상담 후 교사가 기록하면 좋은 내용 1", "기록 내용 2", "기록 내용 3"],
  "next_step_hint": "상담 후 교사 메모를 입력하면 지원 영역 구조화, 지역기관 추천, 회의록 초안 생성으로 연결할 수 있다."
}}
""".strip()


def build_counseling_question_repair_prompt(validation_error: str, previous_output: str) -> str:
    return f"""
이전 출력은 상담 질문 생성 조건을 충족하지 못했습니다.
검증 실패 사유: {validation_error}

아래 조건을 반영하여 다시 JSON 형식으로만 작성하세요.
1. recommended_questions는 5~8개일 것
2. 모든 질문에는 question, purpose, linked_area, based_on, teacher_caution 필드가 있어야 함
3. linked_area는 진로, 학업, 심리정서, 복지경제, 긴급확인 중 하나만 사용할 것
4. 금지 표현, 진단 표현, 취조형 표현, 비교·죄책감 유발 표현, 위협 표현을 사용하지 말 것
5. 학생을 진단하거나 낙인찍지 말 것
6. 공식 체크리스트 문항을 그대로 복사하지 말고 부드러운 개방형 상담 질문으로 바꿀 것

[이전 출력]
{previous_output}
""".strip()


def render_counseling_question_section() -> None:
    first_check_result = st.session_state.get("first_check_result")
    red_flag_result = st.session_state.get("red_flag_result", {"urgent_flag": False, "urgent_flag_items": []})
    active_deep_rules = st.session_state.get("active_deep_rules", [])
    counseling_areas = st.session_state.get("counseling_consideration_areas", [])
    if not first_check_result or not counseling_areas:
        st.info("체크리스트 결과 계산 후 상담 질문을 만들 수 있습니다.")
        return
    official_df, err = load_official_checklist_reference()
    if err:
        st.warning("상담 질문 생성을 위한 참고 자료를 읽지 못했습니다.")
        return
    official_context = get_official_checklist_context(official_df, counseling_areas, red_flag_result)
    st.session_state["official_checklist_context"] = official_context
    activate = bool(first_check_result.get("activate_counseling_form"))
    if activate:
        st.info("현재 결과에서는 2차 상담 질문 생성이 권장됩니다.")
    else:
        st.caption("교사가 필요하다고 판단하면 상담 질문을 생성할 수 있습니다.")
    if not get_gemini_api_key():
        st.warning("Gemini API 키가 설정되어 있지 않아 상담 질문을 생성할 수 없습니다.")
    if st.button("2차 상담 질문 생성하기", type="primary", use_container_width=True, key="btn_generate_questions"):
        if not get_gemini_api_key():
            st.warning("API 키가 설정되지 않아 상담 질문 생성을 실행할 수 없습니다.")
        else:
            clear_pipeline_from("questions_downstream")
            payload = {
                "first_check_result": first_check_result,
                "red_flag_result": red_flag_result,
                "active_deep_rules": active_deep_rules,
                "counseling_consideration_areas": counseling_areas,
                "context_result": st.session_state.get("context_result", {}),
                "official_checklist_context": official_context,
            }
            payload_hash = stable_payload_hash(payload)
            if use_cached_llm_result("generated_counseling_questions", "generated_counseling_questions_payload_hash", payload_hash):
                st.info("같은 입력값으로 이미 생성된 상담 질문을 재사용합니다.")
                result = {"success": True, "data": st.session_state["generated_counseling_questions"], "warnings": []}
            else:
                with st.spinner("2차 상담 질문을 생성하고 있습니다..."):
                    result = call_llm_with_validation(
                        build_counseling_question_system_prompt(),
                        build_counseling_question_user_prompt(payload),
                        validate_counseling_question_output,
                        build_counseling_question_repair_prompt,
                        validation_kwargs={"red_flag_result": red_flag_result, "counseling_consideration_areas": counseling_areas},
                    )
            if result["success"]:
                save_llm_result("generated_counseling_questions", "generated_counseling_questions_payload_hash", payload_hash, result["data"])
                st.success("2차 상담 질문 생성이 완료되었습니다.")
            else:
                st.warning(user_friendly_generation_error(result.get("error"), "상담 질문 생성"))
    data = st.session_state.get("generated_counseling_questions")
    if data:
        summary = data.get("counseling_focus_summary", "")
        if summary:
            st.markdown(f"<div class='callout'>{display_text(summary)}</div>", unsafe_allow_html=True)
        for idx, q in enumerate(data.get("recommended_questions", []), start=1):
            qid = q.get("question_id") or f"Q{idx}"
            with st.container(border=True):
                st.markdown(
                    f"<div class='question-text'>{html.escape(str(qid))}. {html.escape(display_text(q.get('question', '')))}</div>",
                    unsafe_allow_html=True,
                )
                with st.expander("상세보기", expanded=False):
                    detail_rows = [
                        {"항목": "질문 목적", "내용": display_text(q.get("purpose", ""))},
                        {"항목": "연결 영역", "내용": display_area_name(q.get("linked_area", ""))},
                        {"항목": "교사 유의사항", "내용": display_text(q.get("teacher_caution", ""))},
                        {"항목": "추가 확인", "내용": display_text(q.get("follow_up_if_needed", ""))},
                    ]
                    st.dataframe(pd.DataFrame(detail_rows), use_container_width=True, hide_index=True)

# ------------------------------ 상담 결과 분석 ------------------------------
def build_counseling_analysis_payload(
    first_check_result: Dict[str, Any],
    active_deep_rules: List[Dict[str, Any]],
    generated_counseling_questions: Dict[str, Any],
    teacher_counseling_note: str,
    teacher_support_judgment: str,
    existing_support_info: str,
) -> Dict[str, Any]:
    return {
        "first_check_result": first_check_result,
        "active_deep_rules": active_deep_rules,
        "generated_counseling_questions": generated_counseling_questions,
        "counseling_consideration_areas": st.session_state.get("counseling_consideration_areas", []),
        "context_result": st.session_state.get("context_result", {}),
        "official_checklist_context": st.session_state.get("official_checklist_context", []),
        "teacher_counseling_note": teacher_counseling_note,
        "teacher_support_judgment": teacher_support_judgment,
        "existing_support_info": existing_support_info,
    }


def build_counseling_analysis_system_prompt() -> str:
    return """
너는 학생맞춤통합지원 업무를 돕는 교사용 AI 상담 결과 분석 보조도구이다.
교사가 입력한 2차 상담 결과 메모를 읽고 맞춤 검색과 회의록 초안 생성에 사용할 수 있도록 핵심 정보를 구조화한다.
학생을 진단하거나 판정하지 말고, 상담 메모와 제공된 자료에 근거한 정보만 정리한다.
가장 중요한 기준은 교사의 2차 상담 결과 메모이다. 1차 체크리스트와 심층 유도 분석은 참고자료일 뿐이며, 상담 메모에서 직접 확인되지 않은 지원 영역을 primary_area나 key_signals에 넣지 않는다.
상담 메모에서 어떤 영역의 어려움이 없거나 낮다고 명시된 경우, 그 영역은 지원 영역으로 분류하지 않는다.
이 단계에서는 urgent_flag, Red Flag, 긴급확인 분기를 사용하지 않는다.
출력은 반드시 지정된 JSON 형식으로만 작성한다.
""".strip()


def build_counseling_analysis_user_prompt(payload: Dict[str, Any]) -> str:
    return f"""
아래 자료를 바탕으로 2차 상담 결과를 구조화하라.
이 단계의 목적은 기관 추천이나 회의록 완성이 아니라, 다음 단계인 맞춤 검색과 회의록 초안 생성을 위한 구조화된 입력값을 만드는 것이다.
이 단계에서는 urgent_flag, Red Flag, 긴급확인 분기를 사용하지 않는다.

[1차 체크리스트 결과]
{_safe_json(payload.get('first_check_result'))}

[활성화된 심층 유도 분석]
{_safe_json(payload.get('active_deep_rules'))}

[AI가 생성한 2차 상담 질문]
{_safe_json(payload.get('generated_counseling_questions'))}

[교사의 2차 상담 결과 메모]
{payload.get('teacher_counseling_note')}

[교사의 지원 필요 판단]
{payload.get('teacher_support_judgment')}

[기존 지원 여부]
{payload.get('existing_support_info')}

[1차 체크리스트에서 참고할 수 있는 영역]
{_safe_json(payload.get('counseling_consideration_areas'))}

주의:
- 위 영역들은 참고자료일 뿐이다.
- 상담 메모에서 직접 확인된 내용만 primary_area, key_signals, rag_search_queries에 반영하라.
- 상담 메모에 “문제 없음”, “어려움 없음”, “성적이 좋음”, “경제적 부담 없음”처럼 부정 또는 양호함이 적힌 영역은 제외하라.
- 교사의 지원 필요 판단은 반드시 analysis_summary.support_needed에 그대로 반영하라.
- 교사의 지원 필요 판단이 “현재 유지”이면 support_needed도 반드시 “현재 유지”로 작성하고, 기관 검색어는 비워 두거나 최소화하라.
- 예를 들어 학업에는 어려움이 없고 진로 무관심과 친구관계 어려움이 확인되면, primary_area와 linked_areas는 진로·심리정서 중심으로 작성하라.

출력 형식:
{{
  "analysis_summary": {{
    "one_sentence_summary": "상담 결과를 한 문장으로 요약",
    "support_needed": "현재 유지 / 추가 관찰 / 지원 검토 필요 / 판단 보류",
    "support_needed_reason": "지원 필요 판단의 근거",
    "analysis_note": "분석 시 주의할 점 또는 상담 메모의 한계"
  }},
  "primary_area": "진로 / 학업 / 심리정서 / 복지경제 / 공통",
  "key_signals": [
    {{"signal": "주요 신호", "linked_areas": ["심리정서"], "evidence_text": "교사 상담 메모의 실제 표현", "interpretation": "검토 방향"}}
  ],
  "rag_search_queries": [
    {{"query": "맞춤 검색 질의", "target_collection": "policy_chunks / service_catalog / resource_catalog", "purpose": "검색 목적"}}
  ],
  "meeting_record_inputs": {{
    "counseling_summary": "회의록용 상담 요약",
    "discussion_points": ["논의 안건"],
    "teacher_confirmation_items": ["교사 확인사항"],
    "guardian_contact_considerations": ["보호자 상담 또는 동의 관련 확인사항"],
    "follow_up_items": ["후속 관찰 또는 사후관리 항목"]
  }},
  "safety_and_ethics_note": "AI 결과는 자동 판정이 아니라 교사와 학교 협의체 검토를 위한 참고자료입니다."
}}
""".strip()


def build_counseling_analysis_repair_prompt(validation_error: str, previous_output: str) -> str:
    return f"""
이전 출력은 상담 결과 분석 조건을 충족하지 못했습니다.
검증 실패 사유: {validation_error}
반드시 JSON 형식으로만 다시 작성하세요. urgent_flag, urgent_reasons, urgent_notice, 긴급확인 관련 필드는 출력하지 마세요.
지원 검토 필요인 경우 rag_search_queries를 최소 2개 이상 포함하세요. 상담 메모에 없는 사실을 만들지 마세요.

[이전 출력]
{previous_output}
""".strip()


def render_counseling_analysis_section() -> None:
    st.write("학생과 2차 상담을 진행한 뒤, 상담 결과를 간단히 기록하면 지원 검토 방향을 정리합니다.")
    note = st.text_area(
        "2차 상담 결과 메모",
        value=st.session_state.get("teacher_counseling_note", ""),
        placeholder="예: 학생은 수업 시간에 엎드리는 이유가 잠을 잘 못 자서라고 말함. 친구들과 어울리는 것이 부담스럽고 쉬는 시간에는 혼자 있는 것이 편하다고 답함.",
        height=130,
        key="teacher_counseling_note_input",
    )
    col1, col2 = st.columns(2)
    with col1:
        judgment = st.selectbox("교사의 지원 필요 판단", ["현재 유지", "추가 관찰", "지원 검토 필요", "판단 보류"], index=2, key="teacher_support_judgment_input")
    with col2:
        existing = st.text_input("기존 지원 여부", value=st.session_state.get("existing_support_info", "기존 지원 없음"), placeholder="예: 기존 지원 없음 / 현재 Wee클래스 상담 중 / 확인 필요")
    analysis_dirty = counseling_analysis_input_changed(note, judgment, existing)
    if analysis_dirty:
        clear_pipeline_from("analysis")
        st.info("상담 메모 또는 입력값이 변경되었습니다. 다시 분석하면 이후 단계가 새 결과에 맞춰 열립니다.")
    if not get_gemini_api_key():
        st.warning("Gemini API 키가 설정되어 있지 않습니다. Streamlit secrets 또는 환경변수에 GEMINI_API_KEY를 설정해 주세요.")
    if st.button("상담 결과 분석하기", type="primary", use_container_width=True, key="btn_analyze_note"):
        if not note.strip():
            st.warning("상담 결과 메모를 입력해 주세요.")
        elif not st.session_state.get("generated_counseling_questions"):
            st.warning("먼저 2차 상담 질문을 생성해 주세요.")
        elif not st.session_state.get("first_check_result"):
            st.warning("먼저 1차 체크리스트를 완료해 주세요.")
        elif not get_gemini_api_key():
            st.warning("API 키가 설정되지 않아 상담 결과 분석을 실행할 수 없습니다.")
        else:
            clear_pipeline_from("analysis_downstream")
            payload = build_counseling_analysis_payload(
                st.session_state["first_check_result"],
                st.session_state.get("active_deep_rules", []),
                st.session_state["generated_counseling_questions"],
                note,
                judgment,
                existing,
            )
            payload_hash = stable_payload_hash(payload)
            if use_cached_llm_result("structured_counseling_analysis", "structured_counseling_analysis_payload_hash", payload_hash):
                st.info("같은 상담 메모로 이미 분석된 결과를 재사용합니다. API를 다시 호출하지 않았습니다.")
                result = {"success": True, "data": st.session_state["structured_counseling_analysis"], "warnings": []}
            else:
                with st.spinner("상담 결과를 분석하고 있습니다..."):
                    result = call_llm_with_validation(
                        build_counseling_analysis_system_prompt(),
                        build_counseling_analysis_user_prompt(payload),
                        validate_counseling_analysis_output,
                        build_counseling_analysis_repair_prompt,
                        validation_kwargs={"teacher_counseling_note": note},
                    )
            if result["success"]:
                processed_data = postprocess_counseling_analysis_result(result["data"], note)
                save_llm_result("structured_counseling_analysis", "structured_counseling_analysis_payload_hash", payload_hash, processed_data)
                st.session_state["teacher_counseling_note"] = note
                st.session_state["teacher_support_judgment"] = judgment
                st.session_state["existing_support_info"] = existing
                st.success("상담 결과 분석이 완료되었습니다.")
            else:
                st.warning(user_friendly_generation_error(result.get("error"), "상담 결과 분석"))
    data = st.session_state.get("structured_counseling_analysis")
    if data and st.session_state.get("teacher_counseling_note"):
        data = postprocess_counseling_analysis_result(data, st.session_state.get("teacher_counseling_note", ""))
        st.session_state["structured_counseling_analysis"] = data
    if data:
        summ = data.get("analysis_summary", {})
        one_line = summ.get("support_needed_reason") or summ.get("one_sentence_summary") or "상담 결과를 바탕으로 추가 검토가 필요합니다."
        target_areas = derive_integrated_support_areas(data)
        secondary = [a for a in data.get("secondary_consideration_areas", []) if a not in target_areas]
        secondary_help = area_help_text("체크리스트상 함께 고려할 수 있는 영역", secondary, "체크리스트상 추가 고려 영역 없음")
        c1, c2 = st.columns([1.35, 1])
        with c1:
            st.markdown(analysis_summary_card(summ.get("support_needed", "-"), one_line), unsafe_allow_html=True)
        with c2:
            area_value = ", ".join(target_areas) if target_areas else ("현재 추가 확인 영역 없음" if summ.get("support_needed") == "현재 유지" else data.get("primary_area", "-"))
            st.markdown(metric_card("상담에서 확인된 지원 영역", area_value, secondary_help), unsafe_allow_html=True)
        st.caption("AI 결과는 자동 판정이 아니라 교사와 학교 협의체 검토를 돕는 참고자료입니다.")


# ------------------------------ 맞춤 검색 ------------------------------
EMBEDDING_MODEL_NAME = "sentence-transformers/distiluse-base-multilingual-cased-v1"




def stable_payload_hash(payload: Any) -> str:
    """동일 입력값에 대해 API 재호출을 막기 위한 안정적인 해시."""
    try:
        raw = json.dumps(payload, ensure_ascii=False, sort_keys=True, default=str)
    except Exception:
        raw = str(payload)
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def use_cached_llm_result(result_key: str, hash_key: str, payload_hash: str) -> bool:
    return bool(st.session_state.get(result_key)) and st.session_state.get(hash_key) == payload_hash


def save_llm_result(result_key: str, hash_key: str, payload_hash: str, data: Any) -> None:
    st.session_state[result_key] = data
    st.session_state[hash_key] = payload_hash


PIPELINE_STATE_KEYS = {
    "checklist": [
        "first_check_result",
        "red_flag_result",
        "context_result",
        "active_deep_rules",
        "counseling_consideration_areas",
        "last_payload",
        "last_checklist_student",
        "official_checklist_context",
        "generated_counseling_questions",
        "structured_counseling_analysis",
        "rag_search_results",
        "resource_recommendation_explanation",
        "generated_document_json",
        "generated_docx_files",
        "generated_counseling_questions_payload_hash",
        "structured_counseling_analysis_payload_hash",
        "resource_recommendation_explanation_payload_hash",
        "generated_document_json_payload_hash",
    ],
    # 1차 체크리스트를 다시 계산했을 때 2차 상담 이후 결과까지 모두 정리한다.
    "questions": [
        "official_checklist_context",
        "generated_counseling_questions",
        "generated_counseling_questions_payload_hash",
        "structured_counseling_analysis",
        "rag_search_results",
        "resource_recommendation_explanation",
        "generated_document_json",
        "generated_docx_files",
        "structured_counseling_analysis_payload_hash",
        "resource_recommendation_explanation_payload_hash",
        "generated_document_json_payload_hash",
    ],
    # 상담 질문을 다시 생성했을 때는 기존 질문은 재사용할 수 있게 두고 이후 결과만 정리한다.
    "questions_downstream": [
        "structured_counseling_analysis",
        "rag_search_results",
        "resource_recommendation_explanation",
        "generated_document_json",
        "generated_docx_files",
        "structured_counseling_analysis_payload_hash",
        "resource_recommendation_explanation_payload_hash",
        "generated_document_json_payload_hash",
    ],
    # 상담 메모 입력값 자체가 바뀐 경우에는 분석 결과부터 다시 만든다.
    "analysis": [
        "structured_counseling_analysis",
        "rag_search_results",
        "resource_recommendation_explanation",
        "generated_document_json",
        "generated_docx_files",
        "structured_counseling_analysis_payload_hash",
        "resource_recommendation_explanation_payload_hash",
        "generated_document_json_payload_hash",
    ],
    # 상담 결과 분석 버튼을 다시 누른 경우, 분석 결과 캐시는 남겨도 이후 추천/문서는 새로 열리게 한다.
    "analysis_downstream": [
        "rag_search_results",
        "resource_recommendation_explanation",
        "generated_document_json",
        "generated_docx_files",
        "resource_recommendation_explanation_payload_hash",
        "generated_document_json_payload_hash",
    ],
    "recommendation": [
        "rag_search_results",
        "resource_recommendation_explanation",
        "generated_document_json",
        "generated_docx_files",
        "resource_recommendation_explanation_payload_hash",
        "generated_document_json_payload_hash",
    ],
    "document": [
        "generated_document_json",
        "generated_docx_files",
        "generated_document_json_payload_hash",
    ],
}


def clear_pipeline_from(stage: str) -> None:
    """이전 단계가 바뀌었을 때 이후 단계 결과를 정리한다."""
    for key in PIPELINE_STATE_KEYS.get(stage, []):
        st.session_state.pop(key, None)


def normalize_response_values(responses: Dict[str, Any]) -> Dict[str, int]:
    normalized: Dict[str, int] = {}
    for key, value in (responses or {}).items():
        try:
            normalized[str(key)] = int(value)
        except Exception:
            normalized[str(key)] = 0
    return normalized


def checklist_input_changed(selected_student: str, responses: Dict[str, Any]) -> bool:
    stored = st.session_state.get("checklist_responses", {}).get(selected_student)
    if stored is None:
        return False
    return normalize_response_values(stored) != normalize_response_values(responses)




def build_demo_checklist_responses_from_student(student: pd.Series, items_df: pd.DataFrame) -> Dict[str, int]:
    """상세 리포트에서 심층 파악 대상 학생을 바로 상담 질문 생성 단계로 연결하기 위한 기본 체크리스트 응답."""
    responses: Dict[str, int] = {}
    if items_df is None or items_df.empty:
        return responses

    active_df = items_df.copy()
    if "active" in active_df.columns:
        active_df = active_df[active_df["active"].astype(str).str.upper().str.strip() == "Y"]
    for _, row in active_df.iterrows():
        responses[normalize_text(row.get("item_id"))] = 0

    for area in SUPPORT_AREAS:
        try:
            intensity = int(student.get(area, 0))
        except Exception:
            intensity = 0
        if intensity <= 0:
            continue
        area_items = active_df[active_df["domain"].map(normalize_area_name) == area].copy()
        if area_items.empty:
            continue
        # 4단계 점 표시값을 0~2점 체크리스트로 자연스럽게 펼친다.
        remaining = min(max(intensity, 0), len(area_items) * 2)
        for _, row in area_items.iterrows():
            item_id = normalize_text(row.get("item_id"))
            if remaining >= 2:
                responses[item_id] = 2
                remaining -= 2
            elif remaining == 1:
                responses[item_id] = 1
                remaining -= 1
            else:
                break

    # 우선 확인 신호가 있는 데모 학생은 정서-C를 기본 선택한다.
    if bool(student.get("RedFlag", False)):
        for _, row in active_df.iterrows():
            item_id = normalize_text(row.get("item_id"))
            item_code = normalize_text(row.get("item_code"))
            if item_id == "EMO_C" or item_code == "정서-C":
                responses[item_id] = max(1, responses.get(item_id, 0))
                break
    return responses


def calculate_and_store_checklist_for_student(
    selected_student: str,
    responses: Dict[str, int],
    items_df: pd.DataFrame,
    rule_map_df: Optional[pd.DataFrame],
    deep_rules_df: Optional[pd.DataFrame],
    *,
    update_student_row: bool = True,
) -> None:
    """1차 체크리스트 결과 계산과 세션 저장을 한 곳에서 처리한다."""
    clear_pipeline_from("checklist")
    first_result = calculate_checklist_scores(items_df, responses)
    red_flag_result = detect_red_flags(items_df, responses)
    active_deep_rules = activate_deep_rules(responses, rule_map_df, deep_rules_df)
    counseling_areas = derive_counseling_consideration_areas(
        first_result["domain_scores"], responses, red_flag_result, active_deep_rules, items_df
    )
    context_result = calculate_context_result(
        first_result["primary_areas"],
        first_result["student_raw_score"],
        first_result["student_scaled_score"],
        red_flag_result,
        st.session_state.get("selected_school_context"),
        st.session_state.get("selected_region_context"),
    )
    stage_result = {
        "score_based_stage": context_result["score_based_stage"],
        "score_based_action": context_result["score_based_action"],
        "final_action_stage": context_result["final_action_stage"],
        "final_action": context_result["final_action"],
        "final_action_reason": context_result["final_action_reason"],
        "activate_counseling_form": context_result["activate_counseling_form"],
    }
    payload = build_counseling_payload(
        first_result, red_flag_result, context_result, active_deep_rules, counseling_areas, stage_result
    )

    st.session_state["first_check_result"] = payload["first_check_result"]
    st.session_state["red_flag_result"] = payload["red_flag_result"]
    st.session_state["context_result"] = payload["context_result"]
    st.session_state["active_deep_rules"] = payload["active_deep_rules"]
    st.session_state["counseling_consideration_areas"] = counseling_areas
    st.session_state["last_payload"] = payload
    st.session_state["last_checklist_student"] = selected_student
    st.session_state["selected_student_for_checklist"] = selected_student
    st.session_state.checklist_responses[selected_student] = normalize_response_values(responses)

    if update_student_row:
        all_df = st.session_state.students.copy()
        idx = all_df.index[all_df["학생코드"] == selected_student]
        if len(idx) > 0:
            idx0 = idx[0]
            for _, row in first_result["domain_scores"].iterrows():
                area = row["지원 영역"]
                all_df.at[idx0, area] = min(4, int(round(float(row["domain_scaled_score"]) / 25)))
            all_df.at[idx0, "RedFlag"] = bool(red_flag_result.get("urgent_flag"))
            all_df.at[idx0, "최종단계"] = context_result["final_action_stage"]
            all_df.at[idx0, "권장Action"] = context_result["final_action"]
            checked_items: List[str] = []
            for _, row in items_df.iterrows():
                item_id = normalize_text(row.get("item_id"))
                if int(responses.get(item_id, 0)) >= 1:
                    checked_items.append(clean_checklist_item_label(row.get("item_text")))
            all_df.at[idx0, "주요신호"] = ", ".join(checked_items[:3]) if checked_items else "선택된 신호 없음"
            all_df.at[idx0, "기한"] = str(date.today()) if context_result["final_action_stage"] != "일상적 관찰" else "-"
            st.session_state.students = all_df


def prepare_student_checklist_and_open(selected_student: str, student: pd.Series) -> None:
    """학생 상세 리포트에서 1차 체크리스트 탭으로 이동하면서 입력값과 계산 결과를 준비한다."""
    items_df = get_active_items_df()
    if items_df.empty:
        st.warning("체크리스트 자료를 불러올 수 없어 이동할 수 없습니다.")
        return
    rule_map_df = st.session_state.get("rule_map_df")
    deep_rules_df = st.session_state.get("deep_rules_df")
    responses = st.session_state.get("checklist_responses", {}).get(selected_student)
    if not responses:
        responses = build_demo_checklist_responses_from_student(student, items_df)
    calculate_and_store_checklist_for_student(
        selected_student,
        normalize_response_values(responses),
        items_df,
        rule_map_df,
        deep_rules_df,
        update_student_row=True,
    )
    st.session_state["pending_page"] = "1차 체크리스트"
    st.session_state["selected_student_for_checklist"] = selected_student

def counseling_analysis_input_changed(note: str, judgment: str, existing: str) -> bool:
    if not st.session_state.get("structured_counseling_analysis"):
        return False
    return (
        normalize_text(note) != normalize_text(st.session_state.get("teacher_counseling_note", ""))
        or normalize_text(judgment) != normalize_text(st.session_state.get("teacher_support_judgment", ""))
        or normalize_text(existing) != normalize_text(st.session_state.get("existing_support_info", ""))
    )


def parse_pipe_list(value: Any) -> List[str]:
    if value is None:
        return []
    text = str(value).strip()
    if not text or text.lower() == "nan":
        return []
    for sep in [",", ";", "/"]:
        text = text.replace(sep, "|")
    return [x.strip() for x in text.split("|") if x.strip()]


def normalize_school_level(value: Any) -> str:
    text = normalize_text(value)
    if "초" in text:
        return "초"
    if "중" in text:
        return "중"
    if "고" in text:
        return "고"
    return text


def safe_float(value: Any) -> Optional[float]:
    try:
        if value is None or str(value).strip() == "":
            return None
        v = float(value)
        if math.isnan(v):
            return None
        return v
    except Exception:
        return None


def get_allowed_districts(selected_school_district: str, adjacency_map: Dict[str, Any]) -> List[str]:
    if not selected_school_district:
        return []
    if "district_adjacency" in adjacency_map and isinstance(adjacency_map["district_adjacency"], dict):
        adjacency_map = adjacency_map["district_adjacency"]
    return [selected_school_district] + list(adjacency_map.get(selected_school_district, []))


def is_common_or_wide_area_candidate(metadata: Dict[str, Any]) -> bool:
    text = " ".join(str(metadata.get(k, "")) for k in ["district", "filter_region_scope", "region_scope", "resource_scope", "filter_district_list"])
    return any(x in text for x in ["서울공통", "광역", "전체", "서울시 전체", "공통"])


def haversine_km(lat1: float, lon1: float, lat2: float, lon2: float) -> float:
    r = 6371.0
    phi1, phi2 = math.radians(lat1), math.radians(lat2)
    dphi = math.radians(lat2 - lat1)
    dlambda = math.radians(lon2 - lon1)
    a = math.sin(dphi / 2) ** 2 + math.cos(phi1) * math.cos(phi2) * math.sin(dlambda / 2) ** 2
    return 2 * r * math.atan2(math.sqrt(a), math.sqrt(1 - a))


def distance_to_similarity(distance: Any) -> float:
    d = float(distance or 0)
    if 0 <= d <= 1:
        return max(0.0, min(1.0, 1.0 - d))
    return max(0.0, min(1.0, 1.0 / (1.0 + d)))


def ensure_chroma_db() -> Optional[Path]:
    chroma_dir = APP_DIR / "chroma_db"
    if chroma_dir.exists():
        return chroma_dir
    zip_url = None
    try:
        zip_url = st.secrets.get("CHROMA_DB_ZIP_URL", None)
    except Exception:
        zip_url = None
    zip_url = zip_url or os.getenv("CHROMA_DB_ZIP_URL")
    if not zip_url:
        return None
    try:
        import requests
        import zipfile
        tmp_zip = APP_DIR / "chroma_db.zip"
        with st.spinner("지원기관 추천 자료를 준비하고 있습니다..."):
            r = requests.get(zip_url, timeout=120)
            r.raise_for_status()
            tmp_zip.write_bytes(r.content)
            with zipfile.ZipFile(tmp_zip, "r") as zf:
                zf.extractall(APP_DIR)
        if chroma_dir.exists():
            return chroma_dir
    except Exception as exc:
        st.error(f"지원기관 추천 자료 준비에 실패했습니다: {exc}")
        return None
    return chroma_dir if chroma_dir.exists() else None


@st.cache_resource(show_spinner=False)
def get_embedding_model():
    from sentence_transformers import SentenceTransformer
    return SentenceTransformer(EMBEDDING_MODEL_NAME)


@st.cache_resource(show_spinner=False)
def init_chroma_client_cached(chroma_path_str: str):
    import chromadb
    return chromadb.PersistentClient(path=chroma_path_str)


def query_chroma_collection(client: Any, collection_name: str, query: str, n_results: int = 10) -> List[Dict[str, Any]]:
    collection = client.get_collection(collection_name)
    model = get_embedding_model()
    emb = model.encode([query], normalize_embeddings=True).tolist()
    result = collection.query(query_embeddings=emb, n_results=n_results, include=["documents", "metadatas", "distances"])
    out: List[Dict[str, Any]] = []
    ids = result.get("ids", [[]])[0]
    docs = result.get("documents", [[]])[0]
    metas = result.get("metadatas", [[]])[0]
    dists = result.get("distances", [[]])[0]
    for i, doc_id in enumerate(ids):
        meta = metas[i] if i < len(metas) and metas[i] is not None else {}
        out.append({"id": doc_id, "document_text": docs[i] if i < len(docs) else "", "metadata": meta, "distance": dists[i] if i < len(dists) else None, "matched_query": query, "collection": collection_name})
    return out




def resolve_candidate_district(metadata: Dict[str, Any], allowed_districts: Optional[List[str]] = None) -> str:
    """metadata의 district가 잘못 들어간 경우 filter_district_list/address를 함께 보고 실제 자치구를 고른다."""
    allowed_districts = allowed_districts or []
    district = str(metadata.get("district") or metadata.get("자치구") or "").strip()
    dist_list = parse_pipe_list(metadata.get("filter_district_list") or metadata.get("districts_covered") or metadata.get("district") or metadata.get("자치구"))
    if allowed_districts:
        for d in dist_list:
            if d in allowed_districts:
                return d
        for key in ["address", "주소", "resource_address", "기관주소"]:
            text = str(metadata.get(key, ""))
            for d in allowed_districts:
                if d in text:
                    return d
    if district:
        return district
    return dist_list[0] if dist_list else ""


def extract_support_areas_from_metadata(metadata: Dict[str, Any]) -> List[str]:
    """기관 metadata에 담긴 지원영역을 표준 영역 리스트로 변환한다."""
    support_list = [
        normalize_area_name(x)
        for x in parse_pipe_list(
            metadata.get("filter_support_area_list")
            or metadata.get("support_area")
            or metadata.get("linked_area")
            or metadata.get("지원영역")
        )
    ]
    # 중복 제거 및 공백 제거
    out: List[str] = []
    for area in support_list:
        if area and area not in out:
            out.append(area)
    return out


def normalize_target_areas(areas: Any) -> List[str]:
    """맞춤 검색·필터링에서 사용할 복합 지원 영역 리스트를 정리한다."""
    if areas is None:
        return []
    if isinstance(areas, str):
        raw = parse_pipe_list(areas)
    elif isinstance(areas, (list, tuple, set)):
        raw = list(areas)
    else:
        raw = [areas]
    out: List[str] = []
    for area in raw:
        norm = normalize_area_name(area)
        if norm in SUPPORT_AREAS and norm not in out:
            out.append(norm)
    return out


def _note_has_pattern(note: str, patterns: List[str]) -> bool:
    compact = re.sub(r"\s+", "", normalize_text(note))
    raw = normalize_text(note)
    return any(re.search(p, raw) or re.search(p, compact) for p in patterns)


def infer_support_areas_from_counseling_note(note: str) -> List[str]:
    """상담 메모에서 실제로 확인되는 지원 영역만 추출한다.

    1차 체크리스트나 심층분석 영역을 그대로 가져오지 않고, 교사가 입력한 상담 메모에 직접 드러난 신호를 우선한다.
    """
    text = normalize_text(note)
    compact = re.sub(r"\s+", "", text)
    keyword_map = {
        "진로": ["진로", "꿈", "장래", "미래", "목표", "하고싶은", "무엇을하고", "부모님이시키", "시키는대로", "관심이없", "무관심", "직업", "전공"],
        "심리정서": ["무기력", "소극", "친구", "어울", "관계", "외롭", "혼자", "의지", "자존감", "마음", "불편", "부담", "위축", "기분", "정서", "감정"],
        "학업": ["학업", "공부", "성적", "과제", "수업", "숙제", "시험", "집중", "결석", "지각", "학습"],
        "복지경제": ["경제", "금전", "생계", "급식", "결식", "교육비", "기초수급", "차상위", "한부모", "조손", "돌봄", "의복", "위생", "체취", "준비물", "가정형편", "생활비"],
    }
    negative_patterns = {
        "학업": [r"학업.*(문제|어려움).*없", r"공부.*(문제|어려움).*없", r"수업.*(문제|어려움).*없", r"과제.*(문제|어려움).*없", r"성적.*(좋|우수|양호)", r"성적이좋"],
        "복지경제": [r"경제.*(문제|어려움|부담).*없", r"가정.*경제.*(문제|어려움).*없", r"복지.*필요.*없", r"생활.*(문제|어려움).*없", r"돌봄.*(문제|어려움).*없"],
        "심리정서": [r"친구.*(문제|어려움).*없", r"정서.*(문제|어려움).*없", r"마음.*(문제|어려움).*없", r"관계.*(문제|어려움).*없"],
        "진로": [r"진로.*(문제|어려움).*없", r"목표.*(문제|어려움).*없"],
    }
    out: List[str] = []
    for area, keywords in keyword_map.items():
        if any(k in compact for k in keywords):
            out.append(area)
    for area, patterns in negative_patterns.items():
        if area in out and _note_has_pattern(text, patterns):
            out.remove(area)
    return list(dict.fromkeys(out))


def postprocess_counseling_analysis_result(data: Dict[str, Any], teacher_note: str) -> Dict[str, Any]:
    """LLM이 1차 체크리스트 영역을 과도하게 끌고 오는 것을 막고, 상담 메모 근거 영역을 우선한다.

    교사가 선택한 지원 필요 판단은 사용자 입력값이므로 AI 결과보다 우선 적용한다.
    특히 "현재 유지"를 선택한 경우에는 AI가 상담 메모의 단어를 과도하게 해석해
    "지원 검토 필요"로 승격하지 않도록 상태와 후속 검색 기준을 안정적으로 정리한다.
    """
    if not isinstance(data, dict):
        return data
    data = dict(data)
    teacher_judgment = (
        st.session_state.get("teacher_support_judgment_input")
        or st.session_state.get("teacher_support_judgment")
        or ""
    )
    if isinstance(data.get("analysis_summary"), dict) and teacher_judgment in ["현재 유지", "추가 관찰", "지원 검토 필요", "판단 보류"]:
        data["analysis_summary"]["support_needed"] = teacher_judgment
        if teacher_judgment == "현재 유지":
            data["analysis_summary"]["support_needed_reason"] = (
                "교사의 판단에 따라 현재 지원 상태를 유지하며, 필요한 경우 추후 변화 여부를 관찰합니다."
            )
            data["primary_area"] = "공통"
            data["key_signals"] = []
            data["rag_search_queries"] = []
            data["confirmed_support_areas"] = []
            data["secondary_consideration_areas"] = []
            return data
    note_areas = infer_support_areas_from_counseling_note(teacher_note)
    llm_areas: List[str] = []
    primary = normalize_area_name(data.get("primary_area"))
    if primary in SUPPORT_AREAS:
        llm_areas.append(primary)
    for sig in data.get("key_signals", []) or []:
        for area in normalize_target_areas(sig.get("linked_areas", [])):
            if area not in llm_areas:
                llm_areas.append(area)

    # 상담 메모에서 직접 잡힌 영역이 있으면 그 영역을 우선한다.
    confirmed = note_areas or llm_areas
    if note_areas:
        for sig in data.get("key_signals", []) or []:
            linked = [a for a in normalize_target_areas(sig.get("linked_areas", [])) if a in confirmed]
            if not linked:
                # evidence에 직접 키워드가 있는 경우 보정
                inferred = infer_support_areas_from_counseling_note(str(sig.get("evidence_text", "")) + " " + str(sig.get("signal", "")))
                linked = [a for a in inferred if a in confirmed]
            sig["linked_areas"] = linked or confirmed[:1]
        if data.get("primary_area") not in confirmed:
            data["primary_area"] = confirmed[0] if confirmed else "공통"

    # 검색어도 확인된 영역 중심으로 보정한다.
    if confirmed:
        rq = data.get("rag_search_queries", []) or []
        filtered = []
        for q in rq:
            qtext = str(q.get("query", ""))
            if any(a in qtext for a in confirmed) or not any(a in qtext for a in SUPPORT_AREAS):
                filtered.append(q)
        for area in confirmed:
            filtered.append({"query": f"{area} 학생맞춤통합지원 지원 절차", "target_collection": "policy_chunks", "purpose": f"{area} 지원 근거 확인"})
            filtered.append({"query": f"{area} 학생 지원기관 상담 지원", "target_collection": "resource_catalog", "purpose": f"{area} 지원기관 후보 확인"})
        # 중복 제거
        seen = set()
        unique = []
        for q in filtered:
            key = (q.get("target_collection"), q.get("query"))
            if key not in seen:
                seen.add(key)
                unique.append(q)
        data["rag_search_queries"] = unique[:8]

    data["confirmed_support_areas"] = confirmed
    # 체크리스트 기반으로 함께 볼 수 있지만 상담 메모에서는 직접 확인되지 않은 영역
    previous = []
    for item in st.session_state.get("counseling_consideration_areas", []) or []:
        area = normalize_area_name(item.get("area"))
        if area in SUPPORT_AREAS and area not in confirmed and area not in previous:
            previous.append(area)
    data["secondary_consideration_areas"] = previous
    return data


def derive_integrated_support_areas(analysis: Dict[str, Any]) -> List[str]:
    """상담 메모에서 직접 확인된 지원 영역을 우선해 추천 단계의 기준 영역을 정한다.

    이전 버전은 1차 체크리스트와 심층분석 영역까지 모두 합쳐 네 영역이 과도하게 표시될 수 있었다.
    현재 버전은 confirmed_support_areas가 있으면 그것을 우선 사용하고, 없을 때만 상담 분석 결과의 primary/key_signals를 사용한다.
    교사가 "현재 유지"로 판단한 경우에는 기관 추천용 영역을 만들지 않는다.
    """
    summary = analysis.get("analysis_summary", {}) if isinstance(analysis, dict) else {}
    if isinstance(summary, dict) and summary.get("support_needed") == "현재 유지":
        return []
    target: List[str] = []

    def add_many(values: Any) -> None:
        for area in normalize_target_areas(values):
            if area not in target:
                target.append(area)

    add_many(analysis.get("confirmed_support_areas"))
    if target:
        return target

    add_many(analysis.get("primary_area"))
    for sig in analysis.get("key_signals", []) or []:
        add_many(sig.get("linked_areas", []))
    for q in analysis.get("rag_search_queries", []) or []:
        qtext = str(q.get("query", ""))
        for area in SUPPORT_AREAS:
            if area in qtext and area not in target:
                target.append(area)
    if not target:
        # 마지막 보조: 상담 분석 결과가 충분하지 않을 때만 체크리스트 영역을 참고한다.
        for area_obj in st.session_state.get("counseling_consideration_areas", []) or []:
            add_many(area_obj.get("area"))
    return target or ["공통"]


def metadata_support_matches(metadata: Dict[str, Any], target_areas: Any) -> bool:
    support_list = extract_support_areas_from_metadata(metadata)
    areas = normalize_target_areas(target_areas)
    if not areas:
        return bool("공통" in support_list or not support_list)
    return bool(any(area in support_list for area in areas) or "공통" in support_list or not support_list)


def get_support_area_matches(metadata: Dict[str, Any], target_areas: Any) -> List[str]:
    support_list = extract_support_areas_from_metadata(metadata)
    areas = normalize_target_areas(target_areas)
    matches = [area for area in areas if area in support_list]
    if not matches and "공통" in support_list:
        matches = ["공통"]
    return matches


def metadata_level_matches(metadata: Dict[str, Any], selected_school_level: str) -> bool:
    level_list = [normalize_school_level(x) for x in parse_pipe_list(metadata.get("filter_school_level_list") or metadata.get("target_school_levels") or metadata.get("target_school_level"))]
    selected_level = normalize_school_level(selected_school_level)
    return bool(not level_list or "전체" in level_list or "공통" in level_list or selected_level in level_list)


def metadata_district_matches(metadata: Dict[str, Any], allowed_districts: List[str]) -> bool:
    if is_common_or_wide_area_candidate(metadata):
        return True
    if not allowed_districts:
        return True
    dist_list = parse_pipe_list(metadata.get("filter_district_list") or metadata.get("districts_covered") or metadata.get("district") or metadata.get("자치구"))
    if any(d in allowed_districts for d in dist_list):
        return True
    address = " ".join(str(metadata.get(k, "")) for k in ["address", "주소", "resource_address", "기관주소"])
    return any(d in address for d in allowed_districts)


def pseudo_distance_for_metadata_candidate(metadata: Dict[str, Any], primary_area: str) -> float:
    """벡터 검색에서 빠졌지만 metadata 필터로 확실히 맞는 후보의 내부 정렬용 거리값."""
    category = str(metadata.get("resource_category") or metadata.get("filter_resource_category") or "")
    name = str(metadata.get("resource_name") or "")
    service_key = str(metadata.get("filter_service_type_ids") or metadata.get("existing_support_match_key") or "")
    if primary_area == "복지경제":
        if "지역교육복지센터" in category or "교육복지센터" in name:
            return 0.35
        if "행정복지센터" in category or "주민센터" in name:
            return 0.48
        if any(k in category + name + service_key for k in ["가족센터", "아동급식", "청소년방과후", "드림스타트"]):
            return 0.50
    if primary_area == "심리정서" and any(k in category + name for k in ["Wee", "상담", "정신건강"]):
        return 0.38
    if primary_area == "학업" and any(k in category + name for k in ["학습", "학력", "진단"]):
        return 0.40
    if primary_area == "진로" and any(k in category + name for k in ["진로", "꿈드림", "학교밖"]):
        return 0.40
    return 0.58


def get_metadata_matched_resource_candidates(
    client: Any,
    target_areas: Any,
    selected_school_level: str,
    selected_school_district: str,
    allowed_districts: List[str],
    max_candidates: int = 80,
) -> List[Dict[str, Any]]:
    """벡터 검색 상위권에 못 오른 후보라도 metadata상 명확히 맞는 기관을 보강한다.

    복합지원에서는 하나의 primary_area만 보지 않고 target_areas 전체를 통과 기준으로 사용한다.
    예: 학업+심리정서+복지경제가 함께 확인되면 세 영역의 기관 후보를 모두 보강한다.
    """
    target_areas = normalize_target_areas(target_areas)
    try:
        collection = client.get_collection("resource_catalog")
        got = collection.get(include=["documents", "metadatas"])
    except Exception:
        return []
    ids = got.get("ids", []) or []
    docs = got.get("documents", []) or []
    metas = got.get("metadatas", []) or []
    out: List[Dict[str, Any]] = []
    for i, meta in enumerate(metas):
        meta = meta or {}
        if not metadata_support_matches(meta, target_areas):
            continue
        if not metadata_level_matches(meta, selected_school_level):
            continue
        if not metadata_district_matches(meta, allowed_districts):
            continue
        doc_id = ids[i] if i < len(ids) else f"metadata_match_{i}"
        doc = docs[i] if i < len(docs) else ""
        matched_areas = get_support_area_matches(meta, target_areas)
        # 여러 영역과 맞는 경우 가장 낮은 pseudo distance를 사용한다.
        pseudo_distances = [pseudo_distance_for_metadata_candidate(meta, area) for area in (matched_areas or target_areas or ["공통"])]
        out.append({
            "id": doc_id,
            "document_text": doc,
            "metadata": meta,
            "distance": min(pseudo_distances) if pseudo_distances else 0.58,
            "matched_query": f"[metadata 보강] {'|'.join(target_areas)} + {selected_school_district} + {selected_school_level}",
            "collection": "resource_catalog",
            "candidate_source": "metadata_filter_backfill",
            "matched_support_areas": matched_areas,
        })
    def sort_key(item: Dict[str, Any]):
        meta = item.get("metadata", {}) or {}
        district = resolve_candidate_district(meta, allowed_districts)
        category = str(meta.get("resource_category") or "")
        name = str(meta.get("resource_name") or "")
        same = 0 if district == selected_school_district else 1
        cat_priority = 0
        # 복지경제가 target에 포함된 경우 접근 가능한 복지 인프라를 조금 우선한다.
        if "복지경제" in target_areas and ("지역교육복지센터" in category or "교육복지센터" in name):
            cat_priority = -3
        elif "복지경제" in target_areas and ("행정복지센터" in category or "주민센터" in name):
            cat_priority = -2
        elif any(area in target_areas for area in ["심리정서", "학업", "진로"]) and any(k in category + name for k in ["Wee", "상담", "학습", "진로", "꿈드림"]):
            cat_priority = -1
        return (same, cat_priority, float(item.get("distance") or 1.0), str(meta.get("resource_name") or ""))
    return sorted(out, key=sort_key)[:max_candidates]


def merge_resource_candidates(*candidate_lists: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """동일 기관 후보 중 더 좋은 거리값을 가진 항목을 유지한다."""
    merged: Dict[str, Dict[str, Any]] = {}
    for candidates in candidate_lists:
        for item in candidates or []:
            meta = item.get("metadata", {}) or {}
            name = str(meta.get("resource_name") or meta.get("기관명") or item.get("id") or "")
            phone = str(meta.get("phone") or meta.get("전화번호") or meta.get("대표전화") or "")
            address = str(meta.get("address") or meta.get("주소") or "")
            key = f"{name}|{phone or address}"
            if not key.strip("|"):
                key = str(item.get("id"))
            old = merged.get(key)
            if old is None or float(item.get("distance") or 9) < float(old.get("distance") or 9):
                merged[key] = item
    return list(merged.values())

def result_to_simple(item: Dict[str, Any]) -> Dict[str, Any]:
    meta = item.get("metadata", {}) or {}
    text = item.get("document_text", "") or ""
    out = {"id": item.get("id"), "text": text, "text_summary": text[:220], "distance": item.get("distance"), "matched_query": item.get("matched_query")}
    out.update({k: meta.get(k, "") for k in meta.keys()})
    # 자주 쓰는 키 보정
    out["title"] = out.get("title") or out.get("source_title") or out.get("service_type") or out.get("resource_name") or "-"
    out["source_doc"] = out.get("source_doc") or out.get("document_name") or out.get("source") or "-"
    out["support_area"] = normalize_area_name(out.get("support_area") or out.get("linked_area") or out.get("filter_support_area_list") or "")
    return out


def make_fallback_queries(analysis: Dict[str, Any], target_areas: Optional[List[str]] = None) -> List[Dict[str, str]]:
    target_areas = normalize_target_areas(target_areas) or normalize_target_areas(analysis.get("primary_area")) or ["공통"]
    signals = " ".join([str(s.get("signal", "")) for s in analysis.get("key_signals", [])[:3]])
    queries: List[Dict[str, str]] = []
    for area in target_areas:
        queries.extend([
            {"query": f"{area} 학생맞춤통합지원 통합지원 절차 {signals}", "target_collection": "policy_chunks", "purpose": f"{area} 공식 지원 절차 근거 검색"},
            {"query": f"{area} 지원서비스 상담 학습 진로 복지 {signals}", "target_collection": "service_catalog", "purpose": f"{area} 서비스 유형 검색"},
            {"query": f"{area} 상담 지원 지역기관 학생 {signals}", "target_collection": "resource_catalog", "purpose": f"{area} 지역기관 후보 검색"},
        ])
        if area == "복지경제":
            queries.extend([
                {"query": f"복지경제 지역교육복지센터 교육복지센터 사례관리 교육복지안전망 {signals}", "target_collection": "resource_catalog", "purpose": "복지경제 지역교육복지센터 후보 검색"},
                {"query": f"복지경제 행정복지센터 동주민센터 복지급여 긴급지원 교육비 생계비 기초수급 차상위 {signals}", "target_collection": "resource_catalog", "purpose": "복지경제 행정복지센터 후보 검색"},
            ])
    # 중복 제거
    seen = set()
    unique: List[Dict[str, str]] = []
    for q in queries:
        key = (q.get("target_collection"), q.get("query"))
        if key not in seen:
            seen.add(key)
            unique.append(q)
    return unique


def filter_resource_candidates(candidates: List[Dict[str, Any]], target_areas: Any, selected_school_level: str, selected_school_district: str, allowed_districts: List[str], existing_support_info: str) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
    target_areas = normalize_target_areas(target_areas)
    filtered: List[Dict[str, Any]] = []
    excluded: List[Dict[str, str]] = []
    for item in candidates:
        meta = item.get("metadata", {}) or {}
        support_list = extract_support_areas_from_metadata(meta)
        matched_areas = get_support_area_matches(meta, target_areas)
        support_pass = bool(matched_areas or "공통" in support_list or not support_list)
        level_list = [normalize_school_level(x) for x in parse_pipe_list(meta.get("filter_school_level_list") or meta.get("target_school_levels") or meta.get("target_school_level"))]
        selected_level = normalize_school_level(selected_school_level)
        level_pass = bool(not level_list or "전체" in level_list or "공통" in level_list or selected_level in level_list)
        dist_list = parse_pipe_list(meta.get("filter_district_list") or meta.get("districts_covered") or meta.get("district") or meta.get("자치구"))
        district = resolve_candidate_district(meta, allowed_districts)
        district_pass = bool(is_common_or_wide_area_candidate(meta) or not allowed_districts or any(d in allowed_districts for d in dist_list) or district in allowed_districts)
        reasons = []
        if not support_pass:
            reasons.append("지원 영역 불일치")
        if not level_pass:
            reasons.append("학교급 불일치")
        if not district_pass:
            reasons.append("지역 범위 밖")
        item["support_area_filter_pass"] = support_pass
        item["school_level_filter_pass"] = level_pass
        item["district_filter_pass"] = district_pass
        item["matched_support_areas"] = matched_areas
        item["target_support_areas"] = target_areas
        item["filter_pass"] = support_pass and level_pass and district_pass
        item["filter_exclusion_reason"] = ", ".join(reasons)
        name = meta.get("resource_name") or meta.get("기관명") or meta.get("name") or item.get("id")
        category = meta.get("resource_category") or meta.get("service_type") or ""
        if existing_support_info and any(k in existing_support_info for k in ["Wee", "위클래스", "위센터"]) and any(k in str(category) + str(name) for k in ["Wee", "위", "위센터"]):
            item["existing_support_status"] = "기존 지원 점검"
        else:
            item["existing_support_status"] = "신규 연계 후보" if not existing_support_info or "없음" in existing_support_info else "확인 필요"
        if item["filter_pass"]:
            filtered.append(item)
        else:
            excluded.append({"resource_name": str(name), "reason": item["filter_exclusion_reason"]})
    summary: Dict[str, Any] = {"total_candidates": len(candidates), "passed": len(filtered), "excluded": len(excluded), "excluded_samples": excluded[:30], "target_areas": target_areas}
    for reason in ["지원 영역 불일치", "학교급 불일치", "지역 범위 밖"]:
        summary[reason] = sum(1 for x in excluded if reason in x.get("reason", ""))
    return filtered, summary


def get_first_coordinate_value(data: Dict[str, Any], keys: List[str]) -> Optional[float]:
    """좌표 컬럼명이 서로 달라도 최대한 찾아서 float로 변환한다."""
    if not isinstance(data, dict):
        return None
    for key in keys:
        if key in data:
            val = safe_float(data.get(key))
            if val is not None:
                return val
    return None


def format_distance_km(value: Any) -> str:
    val = safe_float(value)
    if val is None:
        return "좌표 정보 없음"
    return f"{val:.2f}km"


def calculate_location_score(candidate: Dict[str, Any], selected_school_district: str, school_lat: Optional[float], school_lon: Optional[float], allowed_districts: List[str]) -> Tuple[float, Optional[float]]:
    meta = candidate.get("metadata", {}) or {}
    district = resolve_candidate_district(meta, allowed_districts)
    lat = get_first_coordinate_value(meta, ["latitude", "위도", "lat", "resource_latitude", "resource_lat", "기관위도", "기관_위도", "시설위도", "시설_위도", "center_latitude", "y", "Y"])
    lon = get_first_coordinate_value(meta, ["longitude", "경도", "lon", "lng", "resource_longitude", "resource_lon", "resource_lng", "기관경도", "기관_경도", "시설경도", "시설_경도", "center_longitude", "x", "X"])
    distance = None
    if school_lat is not None and school_lon is not None and lat is not None and lon is not None:
        distance = haversine_km(school_lat, school_lon, lat, lon)
    if district == selected_school_district:
        if distance is not None and distance <= 3:
            return 20.0, round(distance, 2)
        return 18.0, round(distance, 2) if distance is not None else None
    if district in allowed_districts:
        if distance is not None and distance <= 5:
            return 16.0, round(distance, 2)
        return 14.0, round(distance, 2) if distance is not None else None
    if is_common_or_wide_area_candidate(meta):
        return 10.0, round(distance, 2) if distance is not None else None
    if not district:
        return 8.0, round(distance, 2) if distance is not None else None
    return 5.0, round(distance, 2) if distance is not None else None


def rank_resource_candidates(filtered_candidates: List[Dict[str, Any]], selected_school_info: Dict[str, Any], allowed_districts: List[str]) -> List[Dict[str, Any]]:
    school_lat = get_first_coordinate_value(selected_school_info, ["위도", "latitude", "lat", "school_latitude", "학교위도", "y", "Y"])
    school_lon = get_first_coordinate_value(selected_school_info, ["경도", "longitude", "lon", "lng", "school_longitude", "학교경도", "x", "X"])
    district = selected_school_info.get("자치구", "")
    dedup: Dict[str, Dict[str, Any]] = {}
    for item in filtered_candidates:
        meta = item.get("metadata", {}) or {}
        name = meta.get("resource_name") or meta.get("기관명") or meta.get("name") or item.get("id")
        phone = meta.get("phone") or meta.get("대표전화") or meta.get("전화번호") or ""
        address = meta.get("address") or meta.get("주소") or ""
        key = f"{name}|{phone or address}"
        student_fit = distance_to_similarity(item.get("distance")) * 80
        location_score, distance_km = calculate_location_score(item, district, school_lat, school_lon, allowed_districts)
        score = student_fit + location_score
        candidate = {
            "resource_name": name,
            "resource_category": meta.get("resource_category") or meta.get("service_type") or "-",
            "support_area": "|".join(item.get("matched_support_areas") or extract_support_areas_from_metadata(meta) or ["공통"]),
            "matched_support_areas": item.get("matched_support_areas", []),
            "district": resolve_candidate_district(meta, allowed_districts),
            "education_office": meta.get("education_office") or meta.get("교육지원청") or "",
            "address": address,
            "phone": phone,
            "homepage": meta.get("homepage") or meta.get("homepage_url") or meta.get("홈페이지") or "",
            "latitude": get_first_coordinate_value(meta, ["latitude", "위도", "lat", "resource_latitude", "resource_lat", "기관위도", "기관_위도", "시설위도", "시설_위도", "center_latitude", "y", "Y"]),
            "longitude": get_first_coordinate_value(meta, ["longitude", "경도", "lon", "lng", "resource_longitude", "resource_lon", "resource_lng", "기관경도", "기관_경도", "시설경도", "시설_경도", "center_longitude", "x", "X"]),
            "distance_km": distance_km,
            "student_fit_score": round(student_fit, 1),
            "location_score": round(location_score, 1),
            "recommendation_score": round(score, 1),
            "recommendation_fit": "높음" if score >= 85 else ("보통" if score >= 70 else ("보조 검토" if score >= 55 else "낮음")),
            "existing_support_status": item.get("existing_support_status", "확인 필요"),
            "score_breakdown": {"student_fit_score": round(student_fit, 1), "location_score": round(location_score, 1)},
            "matched_query": item.get("matched_query", ""),
            "document_text": item.get("document_text", ""),
            "metadata": meta,
        }
        if key not in dedup or candidate["student_fit_score"] > dedup[key]["student_fit_score"]:
            dedup[key] = candidate
    ranked = sorted(dedup.values(), key=lambda x: x["recommendation_score"], reverse=True)
    for i, item in enumerate(ranked, start=1):
        item["rank"] = i
    return ranked[:5]


def run_rag_search() -> Optional[Dict[str, Any]]:
    analysis = st.session_state.get("structured_counseling_analysis")
    if not analysis:
        st.warning("먼저 2차 상담 결과 분석을 실행해 주세요.")
        return None
    school = st.session_state.get("selected_school_info", {})
    district = school.get("자치구", "")
    level = school.get("학교급", "")
    if not district:
        st.warning("학교 자치구 정보가 없어 지역 필터를 적용할 수 없습니다.")
    chroma_dir = ensure_chroma_db()
    if chroma_dir is None or not chroma_dir.exists():
        st.error("지원기관 추천 자료를 찾을 수 없습니다. 관리자에게 문의해 주세요.")
        return None
    try:
        client = init_chroma_client_cached(str(chroma_dir))
    except Exception as exc:
        st.error(f"검색 자료 연결에 실패했습니다: {exc}")
        return None
    adjacency = load_json_with_fallback(JSON_PATHS.get("district_adjacency", []))
    allowed_districts = get_allowed_districts(district, adjacency)
    existing = st.session_state.get("existing_support_info", "기존 지원 없음")
    target_areas = derive_integrated_support_areas(analysis)
    queries = analysis.get("rag_search_queries") or []
    fallback_queries = make_fallback_queries(analysis, target_areas)
    if not queries:
        queries = list(fallback_queries)
    # LLM이 한 영역 위주 검색어만 만들더라도 복합 영역별 fallback query를 함께 추가한다.
    existing_query_keys = {(q.get("target_collection"), q.get("query")) for q in queries}
    for fq in fallback_queries:
        key = (fq.get("target_collection"), fq.get("query"))
        if key not in existing_query_keys:
            queries.append(fq)
            existing_query_keys.add(key)

    policy_items: List[Dict[str, Any]] = []
    service_items: List[Dict[str, Any]] = []
    resource_raw: List[Dict[str, Any]] = []
    for q in queries:
        collection = q.get("target_collection")
        query = q.get("query")
        try:
            if collection == "policy_chunks":
                policy_items.extend([result_to_simple(x) for x in query_chroma_collection(client, "policy_chunks", query, n_results=6)])
            elif collection == "service_catalog":
                service_items.extend([result_to_simple(x) for x in query_chroma_collection(client, "service_catalog", query, n_results=6)])
            elif collection == "resource_catalog":
                resource_raw.extend(query_chroma_collection(client, "resource_catalog", query, n_results=50))
        except Exception as exc:
            st.warning(f"{collection} 검색 중 오류: {exc}")
    primary_area = normalize_area_name(analysis.get("primary_area", "공통")) or (target_areas[0] if target_areas else "공통")
    metadata_backfill = get_metadata_matched_resource_candidates(client, target_areas, level, district, allowed_districts, max_candidates=160)
    resource_raw_merged = merge_resource_candidates(resource_raw, metadata_backfill)
    filtered, debug = filter_resource_candidates(resource_raw_merged, target_areas, level, district, allowed_districts, existing)
    debug["vector_candidate_count"] = len(resource_raw)
    debug["metadata_backfill_candidate_count"] = len(metadata_backfill)
    debug["merged_candidate_count"] = len(resource_raw_merged)
    ranked = rank_resource_candidates(filtered, school, allowed_districts)
    results = {
        "policy_evidence": policy_items[:8],
        "service_catalog_results": service_items[:8],
        "ranked_resources": ranked,
        "filter_debug_summary": debug,
        "search_context": {
            "primary_area": primary_area,
            "target_areas": target_areas,
            "integrated_support_mode": True,
            "selected_school_district": district,
            "selected_school_level": normalize_school_level(level),
            "selected_school_latitude": get_first_coordinate_value(school, ["위도", "latitude", "lat", "school_latitude", "학교위도", "학교_위도", "y", "Y"]),
            "selected_school_longitude": get_first_coordinate_value(school, ["경도", "longitude", "lon", "lng", "school_longitude", "학교경도", "학교_경도", "x", "X"]),
            "school_coordinate_found": get_first_coordinate_value(school, ["위도", "latitude", "lat", "school_latitude", "학교위도", "학교_위도", "y", "Y"]) is not None and get_first_coordinate_value(school, ["경도", "longitude", "lon", "lng", "school_longitude", "학교경도", "학교_경도", "x", "X"]) is not None,
            "allowed_districts": allowed_districts,
            "existing_support_info": existing,
            "used_query_count": len(queries),
            "vector_candidate_count": len(resource_raw),
            "metadata_backfill_candidate_count": len(metadata_backfill),
            "merged_candidate_count": len(resource_raw_merged),
        },
    }
    st.session_state["rag_search_results"] = results
    return results



def generate_resource_recommendation_for_results(rag: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    if not rag or not st.session_state.get("structured_counseling_analysis"):
        return None
    if not rag.get("ranked_resources"):
        st.session_state["resource_recommendation_explanation"] = {
            "recommendation_summary": {"one_sentence_summary": "조건에 맞는 지역기관 후보가 부족합니다."},
            "recommended_resources": [],
            "if_no_suitable_resource": {"no_resource_flag": True, "reason": "조건에 맞는 지역기관 후보가 부족합니다.", "suggested_next_steps": []},
        }
        return st.session_state["resource_recommendation_explanation"]
    if not get_gemini_api_key():
        st.warning("API 키가 설정되지 않아 기관 추천 이유를 생성하지 못했습니다.")
        return None
    payload = build_resource_recommendation_payload(
        st.session_state.get("structured_counseling_analysis", {}),
        st.session_state.get("first_check_result", {}),
        st.session_state.get("context_result", {}),
        st.session_state.get("existing_support_info", "기존 지원 없음"),
        rag,
    )
    payload_hash = stable_payload_hash(payload)
    if use_cached_llm_result("resource_recommendation_explanation", "resource_recommendation_explanation_payload_hash", payload_hash):
        return st.session_state.get("resource_recommendation_explanation")
    result = call_llm_with_validation(
        build_resource_recommendation_system_prompt(),
        build_resource_recommendation_user_prompt(payload),
        validate_resource_recommendation_output,
        build_resource_recommendation_repair_prompt,
        validation_kwargs={"ranked_resources": rag.get("ranked_resources", []), "policy_evidence": rag.get("policy_evidence", [])},
    )
    if result.get("success"):
        save_llm_result("resource_recommendation_explanation", "resource_recommendation_explanation_payload_hash", payload_hash, result["data"])
        return result["data"]
    st.warning(user_friendly_generation_error(result.get("error"), "지원기관 추천 이유 생성") + " 기관 후보는 먼저 표시합니다.")
    return None


def build_recommendation_reason_map() -> Dict[str, str]:
    data = st.session_state.get("resource_recommendation_explanation") or {}
    out: Dict[str, str] = {}
    for r in data.get("recommended_resources", []) or []:
        name = normalize_text(r.get("resource_name"))
        reasons = r.get("recommendation_reasons", []) or []
        if name:
            out[name] = " / ".join(map(str, reasons)) if isinstance(reasons, list) else str(reasons)
    return out


def _resource_reason_map() -> Dict[str, str]:
    data = st.session_state.get("resource_recommendation_explanation", {}) or {}
    mapping: Dict[str, str] = {}
    for item in data.get("recommended_resources", []) or []:
        name = normalize_text(item.get("resource_name"))
        if not name:
            continue
        reasons = item.get("recommendation_reasons", []) or []
        if isinstance(reasons, list):
            reason_text = " ".join([str(x) for x in reasons[:2] if x])
        else:
            reason_text = str(reasons)
        mapping[re.sub(r"\s+", "", name)] = reason_text
    return mapping


def render_rag_search_section() -> None:
    st.write("상담 결과 분석을 바탕으로 학교 상황과 지역 여건에 맞는 지원기관 후보를 추천합니다.")
    analysis = st.session_state.get("structured_counseling_analysis")
    if not analysis:
        st.info("2차 상담 결과 분석이 완료되면 지원기관을 추천할 수 있습니다.")
        return
    school = st.session_state.get("selected_school_info", {})
    adjacency = load_json_with_fallback(JSON_PATHS.get("district_adjacency", []))
    allowed = get_allowed_districts(school.get("자치구", ""), adjacency)
    target_areas_preview = derive_integrated_support_areas(analysis)
    secondary_preview = [a for a in analysis.get("secondary_consideration_areas", []) if a not in target_areas_preview]
    support_status = (analysis.get("analysis_summary", {}) or {}).get("support_needed", "") if isinstance(analysis, dict) else ""
    if support_status == "현재 유지":
        st.info("상담 결과가 현재 유지로 정리되어 별도 지원기관 추천은 진행하지 않습니다. 필요 시 상담 메모를 수정한 뒤 다시 분석해 주세요.")
        return
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        preview_value = ", ".join(target_areas_preview) if target_areas_preview else analysis.get("primary_area", "-")
        st.markdown(metric_card("상담에서 확인된 영역", preview_value, area_help_text("체크리스트상 함께 고려할 수 있는 영역", secondary_preview, "체크리스트상 추가 고려 영역 없음")), unsafe_allow_html=True)
    with c2:
        st.markdown(metric_card("학교 자치구", school.get("자치구", "-"), "지역 기준"), unsafe_allow_html=True)
    with c3:
        st.markdown(metric_card("학교급", school.get("학교급", "-"), "대상 기준"), unsafe_allow_html=True)
    with c4:
        st.markdown(metric_card("기존 지원", st.session_state.get("existing_support_info", "-"), "중복 지원 점검"), unsafe_allow_html=True)

    if st.button("지원기관 추천하기", type="primary", use_container_width=True, key="btn_run_rag"):
        with st.spinner("지원기관 후보를 찾고 추천 이유를 생성하고 있습니다..."):
            clear_pipeline_from("recommendation")
            results = run_rag_search()
            if results and get_gemini_api_key():
                payload = build_resource_recommendation_payload(
                    st.session_state.get("structured_counseling_analysis", {}),
                    st.session_state.get("first_check_result", {}),
                    st.session_state.get("context_result", {}),
                    st.session_state.get("existing_support_info", ""),
                    results,
                )
                payload_hash = stable_payload_hash(payload)
                if use_cached_llm_result("resource_recommendation_explanation", "resource_recommendation_explanation_payload_hash", payload_hash):
                    pass
                else:
                    st.session_state.pop("resource_recommendation_explanation", None)
                    st.session_state.pop("resource_recommendation_explanation_payload_hash", None)
                    ranked = results.get("ranked_resources", [])
                    policy = results.get("policy_evidence", [])
                    rec_result = call_llm_with_validation(
                        build_resource_recommendation_system_prompt(),
                        build_resource_recommendation_user_prompt(payload),
                        validate_resource_recommendation_output,
                        build_resource_recommendation_repair_prompt,
                        validation_kwargs={"ranked_resources": ranked, "policy_evidence": policy},
                    )
                    if rec_result["success"]:
                        save_llm_result("resource_recommendation_explanation", "resource_recommendation_explanation_payload_hash", payload_hash, rec_result["data"])
                    else:
                        st.warning(user_friendly_generation_error(rec_result.get("error"), "추천 이유 생성") + " 기관 후보는 먼저 표시합니다.")
            elif results and not get_gemini_api_key():
                st.warning("API 키가 설정되어 있지 않아 추천 이유는 표시하지 않습니다.")

    results = st.session_state.get("rag_search_results")
    if results:
        resources = results.get("ranked_resources", [])
        if not resources:
            st.warning("현재 조건에서 바로 제시할 수 있는 기관 후보가 부족합니다. 보조 영역 검토 또는 지역 범위 확장이 필요할 수 있습니다.")
        else:
            reason_map = _resource_reason_map()
            for r in resources:
                name = normalize_text(r.get("resource_name"))
                support_area = ", ".join(r.get("matched_support_areas", []) or normalize_area_list(r.get("support_area"))) or r.get("support_area")
                reason = reason_map.get(re.sub(r"\s+", "", name), "")
                with st.container(border=True):
                    rank = r.get("rank", "-")
                    st.markdown(f"<div class='resource-title'>{html.escape(str(rank))}순위. {html.escape(str(name))}</div>", unsafe_allow_html=True)
                    rows = [
                        ("기관유형", r.get("resource_category")),
                        ("지원 영역", support_area),
                        ("자치구", r.get("district")),
                        ("거리", format_distance_km(r.get("distance_km"))),
                        ("주소", r.get("address")),
                        ("전화번호", r.get("phone")),
                        ("홈페이지", r.get("homepage")),
                        ("기존 지원 상태", r.get("existing_support_status")),
                        ("추천 이유", reason or "상담 결과와 학교·지역 여건을 바탕으로 검토할 수 있는 지원기관입니다."),
                    ]
                    st.markdown(resource_detail_table_html(rows), unsafe_allow_html=True)
            st.caption("추천기관은 교사와 학교 협의체가 검토할 후보입니다. 실제 연계 전 기관 운영 여부와 보호자 동의 여부를 확인해 주세요.")

# ------------------------------ 기관 추천 이유 생성 ------------------------------
def build_resource_recommendation_payload(structured_counseling_analysis: Dict[str, Any], first_check_result: Dict[str, Any], context_result: Dict[str, Any], existing_support_info: str, rag_search_results: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "structured_counseling_analysis": structured_counseling_analysis,
        "first_check_result": first_check_result,
        "context_result": context_result,
        "existing_support_info": existing_support_info,
        "policy_evidence": rag_search_results.get("policy_evidence", []),
        "service_catalog_results": rag_search_results.get("service_catalog_results", []),
        "ranked_resources": rag_search_results.get("ranked_resources", []),
        "search_context": rag_search_results.get("search_context", {}),
        "target_areas": rag_search_results.get("search_context", {}).get("target_areas", []),
    }


def build_resource_recommendation_system_prompt() -> str:
    return """
너는 학생맞춤통합지원 업무를 돕는 교사용 AI 기관 추천 설명 보조도구이다.
제공된 공식 근거, 서비스 유형, 지역기관 후보 안에서만 교사가 검토할 수 있는 기관 추천 설명을 작성한다.
새로운 기관명, 제도, 연락처, 주소를 만들지 않는다. Python 코드가 정한 ranked_resources 순서를 바꾸지 않는다.
학생을 진단하거나 판정하지 않고, 추천은 교사와 학교 협의체가 검토할 후보로 표현한다.
이 단계에서는 Red Flag, urgent_flag, 긴급확인 관련 내용을 사용하지 않는다.
출력은 반드시 지정된 JSON 형식으로 작성한다.
""".strip()


def build_resource_recommendation_user_prompt(payload: Dict[str, Any]) -> str:
    return f"""
아래 자료를 바탕으로 학생맞춤통합지원 지역기관 추천 설명을 생성하라.
기관을 새로 찾지 말고 이미 맞춤 검색과 Python 필터링·순위화를 통해 추려진 후보를 설명하라.
ranked_resources 입력 순서를 반드시 유지하라.
recommended_resources의 rank와 resource_name은 ranked_resources에 있는 값을 글자 하나 바꾸지 말고 그대로 복사하라.
기관 설명을 3개만 작성한다면 반드시 ranked_resources의 1순위, 2순위, 3순위 순서대로 작성하라.
거리, 추천점수, 주소, 전화번호, 홈페이지는 새로 계산하거나 바꾸지 말고 ranked_resources 값을 그대로 사용하라.

[상담 결과 구조화 정보]
{_safe_json(payload.get('structured_counseling_analysis'))}

[1차 체크리스트 결과]
{_safe_json(payload.get('first_check_result'))}

[학교·지역 맥락 참고 정보]
{_safe_json(payload.get('context_result'))}

[기존 지원 여부]
{payload.get('existing_support_info')}

[복합 검토 지원 영역]
{_safe_json(payload.get('target_areas'))}

[맞춤 검색 결과: 공식 근거]
{_safe_json(payload.get('policy_evidence'))}

[맞춤 검색 결과: 서비스 카탈로그 후보]
{_safe_json(payload.get('service_catalog_results'))}

[맞춤 검색 결과: 지역기관 후보]
{_safe_json(payload.get('ranked_resources'))}

[공식 근거 작성 규칙]
- official_basis는 반드시 위 [맞춤 검색 결과: 공식 근거]에 있는 policy_evidence에서만 작성한다.
- official_basis.source_doc에는 policy_evidence의 source_doc 값을 그대로 복사한다.
- service_catalog, resource_catalog, 기관 후보명, 서비스 카탈로그명은 official_basis.source_doc에 쓰지 않는다.
- 서비스 카탈로그와 기관 후보에서 얻은 내용은 recommendation_reasons, teacher_confirmation_items, location_or_access_basis에만 반영한다.
- 사용할 수 있는 policy_evidence가 없으면 official_basis는 빈 리스트 []로 둔다.

출력 형식:
{{
  "recommendation_summary": {{
    "one_sentence_summary": "추천 결과를 한 문장으로 요약",
    "primary_area": "진로 / 학업 / 심리정서 / 복지경제 / 공통",
    "support_needed_status": "현재 유지 / 추가 관찰 / 지원 검토 필요 / 판단 보류",
    "recommendation_mode": "일반 추천 / 기존 지원 점검 / 후보 부족",
    "summary_for_teacher": "교사용 요약 설명"
  }},
  "recommended_resources": [
    {{
      "rank": 1,
      "service_type": "추천서비스 유형",
      "resource_name": "기관명",
      "resource_category": "기관유형",
      "linked_area": "진로 / 학업 / 심리정서 / 복지경제 / 공통 중 하나",
      "district": "자치구",
      "education_office": "교육지원청",
      "address": "주소",
      "phone": "전화번호",
      "homepage": "홈페이지",
      "distance_km": null,
      "recommendation_fit": "높음 / 보통 / 보조 검토",
      "recommendation_score": null,
      "score_breakdown": {{"student_fit_score": null, "location_score": null}},
      "existing_support_status": "신규 연계 후보 / 기존 지원 점검 / 확인 필요",
      "recommendation_reasons": ["추천 이유"],
      "student_context_basis": ["상담 메모 또는 체크리스트와 연결되는 근거"],
      "official_basis": [{{"basis_title": "공식 근거 제목", "basis_summary": "요약", "source_doc": "출처 문서명", "source_page": "출처 페이지"}}],
      "location_or_access_basis": "지역·접근성 관련 설명",
      "teacher_confirmation_items": ["교사가 확인해야 할 사항"],
      "referral_cautions": ["연계 시 유의사항"],
      "meeting_record_sentence": "회의록에 넣을 수 있는 문장"
    }}
  ],
  "if_no_suitable_resource": {{"no_resource_flag": false, "reason": "", "suggested_next_steps": []}},
  "overall_teacher_checklist": ["추천 전 확인사항 1", "확인사항 2", "확인사항 3"],
  "rag_trace_summary": {{"used_policy_chunks": [], "used_service_catalog_items": [], "used_resource_candidates": [], "note": "제공된 맞춤 검색 결과 안에서만 작성"}},
  "safety_and_ethics_note": "AI 추천은 자동 결정이 아니라 교사와 학교 협의체 검토를 위한 참고자료입니다."
}}
""".strip()


def build_resource_recommendation_repair_prompt(validation_error: str, previous_output: str) -> str:
    return f"""
이전 출력은 맞춤 지원기관 추천 설명 조건을 충족하지 못했습니다.
검증 실패 사유: {validation_error}
반드시 JSON 형식으로만 다시 작성하세요.
ranked_resources에 제공된 기관 후보 안에서만 작성하고, 기관명·주소·전화번호·홈페이지를 새로 만들지 말고, 추천 순서를 바꾸지 마세요.
Red Flag, urgent_flag, urgent_notice, 긴급확인 관련 내용은 출력하지 마세요.

[이전 출력]
{previous_output}
""".strip()


def render_resource_recommendation_section() -> None:
    # 기관 추천 이유는 맞춤 지원기관 검색 버튼을 누를 때 자동으로 생성되어 기관 표에 함께 표시됩니다.
    return


# ------------------------------ 문서 생성 ------------------------------
def build_document_generation_payload(structured_counseling_analysis: Dict[str, Any], resource_recommendation_explanation: Dict[str, Any], rag_search_results: Dict[str, Any], first_check_result: Dict[str, Any], document_type: str) -> Dict[str, Any]:
    return {
        "document_type": document_type,
        "first_check_result": first_check_result,
        "structured_counseling_analysis": structured_counseling_analysis,
        "resource_recommendation_explanation": resource_recommendation_explanation,
        "policy_evidence": rag_search_results.get("policy_evidence", []),
        "recommended_resources": resource_recommendation_explanation.get("recommended_resources", []),
    }

def build_document_generation_system_prompt() -> str:
    return """
너는 학생맞춤통합지원 업무를 돕는 회의록 초안 작성 보조도구이다.
제공된 1차 체크리스트 결과, 상담 결과 분석, 지원기관 추천 설명, 공식 근거를 바탕으로 통합지원팀 회의록에 들어갈 서술형 내용을 작성한다.
학생을 진단하거나 판정하지 않는다. 교사의 최종 판단을 대체하지 않는다.
제공된 자료에 없는 사실, 기관명, 연락처, 주소, 제도명을 새로 만들지 않는다.
개인정보를 생성하거나 추정하지 않는다. Red Flag, urgent_flag, 긴급확인 관련 표현은 사용하지 않는다.
회의록 문체는 명사형 종결어미로 작성한다. 예: "확인함", "검토함", "연계 예정임", "추가 관찰 필요함".
"~합니다", "~했습니다", "~것입니다"처럼 설명문 형식의 종결은 사용하지 않는다.
출력은 반드시 지정된 JSON 형식으로만 작성한다.
""".strip()

def build_document_generation_user_prompt(payload: Dict[str, Any]) -> str:
    return f"""
아래 자료를 바탕으로 통합지원팀 회의록에 들어갈 서술형 내용을 JSON으로 작성하라.
개인정보는 제공하지 않는다. 학생명, 생년월일, 연락처, 주소는 생성하지 않는다.
기관명은 제공된 추천기관 안에서만 사용한다.
회의록에 들어갈 문장은 명사형 종결어미로 마무리한다.
예: "상담 메모에서 진로 관심 저하와 또래관계 소극성이 확인됨", "위클래스 상담 경과 확인 후 추가 연계 검토함".
"~합니다", "~했습니다", "~필요합니다" 형식은 사용하지 말고 "~함", "~확인됨", "~필요함"으로 작성한다.

[문서 생성 대상]
{payload.get('document_type')}

[1차 체크리스트 결과]
{_safe_json(payload.get('first_check_result'))}

[상담 결과 분석]
{_safe_json(payload.get('structured_counseling_analysis'))}

[지원기관 추천 설명]
{_safe_json(payload.get('resource_recommendation_explanation'))}

[공식 근거 및 추천기관]
{_safe_json({'policy_evidence': payload.get('policy_evidence'), 'recommended_resources': payload.get('recommended_resources')})}

출력 JSON 형식:
{{
  "meeting_record": {{
    "agenda": "회의 안건 초안",
    "meeting_content": "대상 학생 협의 내용 초안",
    "support_plan": [
      "지원계획 1",
      "지원계획 2"
    ],
    "decision_items": []
  }},
  "safety_and_ethics_note": "AI 결과는 자동 판정이 아니라 교사와 학교 협의체 검토를 위한 초안입니다."
}}

주의:
- support_plan에는 기존 버전의 결정사항에 들어가던 지원 실행 계획을 작성한다.
- decision_items는 반드시 빈 리스트 []로 둔다. 최종 결정사항은 회의 후 교사가 직접 작성한다.
- agenda, meeting_content, support_plan의 각 문장은 회의록 문체에 맞게 명사형 종결어미로 끝낸다.
""".strip()

def build_document_generation_repair_prompt(validation_error: str, previous_output: str) -> str:
    return f"""
이전 출력은 회의록 생성 조건을 충족하지 못했습니다.
검증 실패 사유: {validation_error}
아래 조건을 반드시 반영하여 다시 작성하세요.
1. 반드시 JSON 형식으로만 출력할 것
2. meeting_record를 포함할 것
3. meeting_record에는 agenda, meeting_content, support_plan, decision_items를 포함할 것
4. decision_items는 빈 리스트 []로 둘 것
5. 기관명은 제공된 추천기관 목록 안에서만 사용할 것
6. 개인정보를 생성하거나 추정하지 말 것
7. 학생을 진단하거나 낙인찍는 표현을 쓰지 말 것
8. Red Flag, urgent_flag, 긴급확인 관련 내용은 출력하지 말 것
9. 회의록 문장은 명사형 종결어미로 끝낼 것. 예: 확인함, 검토함, 연계 예정임, 추가 관찰 필요함
10. ~합니다, ~했습니다, ~필요합니다 형식은 사용하지 말 것

[이전 출력]
{previous_output}
""".strip()


def validate_meeting_generation_output(output_text: str, allowed_resource_names: Optional[List[str]] = None) -> Dict[str, Any]:
    try:
        data = json.loads(strip_json_code_fence_local(output_text))
    except Exception as exc:
        return {"ok": False, "message": f"JSON 형식이 올바르지 않습니다: {exc}", "parsed_data": None, "warnings": []}
    if not isinstance(data, dict):
        return {"ok": False, "message": "출력은 JSON 객체여야 합니다.", "parsed_data": None, "warnings": []}
    data = sanitize_llm_parsed_data(data)
    mr = data.get("meeting_record")
    if not isinstance(mr, dict):
        return {"ok": False, "message": "meeting_record가 필요합니다.", "parsed_data": None, "warnings": []}
    for field in ["agenda", "meeting_content", "support_plan", "decision_items"]:
        if field not in mr:
            return {"ok": False, "message": f"meeting_record.{field} 필드가 필요합니다.", "parsed_data": None, "warnings": []}
    if not isinstance(mr.get("support_plan"), list):
        return {"ok": False, "message": "support_plan은 리스트여야 합니다.", "parsed_data": None, "warnings": []}
    if mr.get("decision_items") not in ([], None):
        return {"ok": False, "message": "decision_items는 회의 후 교사가 작성하므로 빈 리스트여야 합니다.", "parsed_data": None, "warnings": []}
    text = _safe_json(data)
    banned = [w for w in HARD_BANNED_COMMON if w in text]
    if banned:
        return {"ok": False, "message": "금지 표현이 포함되었습니다: " + ", ".join(banned), "parsed_data": None, "warnings": []}
    data.setdefault("safety_and_ethics_note", "AI 결과는 자동 판정이 아니라 교사와 학교 협의체 검토를 위한 초안입니다.")
    mr["decision_items"] = []
    return {"ok": True, "message": "검증 통과", "parsed_data": data, "warnings": []}


def strip_json_code_fence_local(text: str) -> str:
    t = str(text or "").strip()
    m = re.match(r"^```(?:json)?\s*(.*?)\s*```$", t, flags=re.IGNORECASE | re.DOTALL)
    return m.group(1).strip() if m else t


def _nominalize_korean_sentence(sentence: str) -> str:
    """회의록용 문장을 가능한 범위에서 명사형 종결로 정리한다."""
    text = display_text(sentence).strip()
    if not text:
        return text
    prefix = ""
    m = re.match(r"^(\s*[-•·]\s*)(.*)$", text)
    if m:
        prefix, text = m.group(1), m.group(2).strip()
    end_punct = "." if text.endswith(".") else ""
    text_core = text[:-1].strip() if end_punct else text

    replacements = [
        ("필요합니다", "필요함"),
        ("필요합니다", "필요함"),
        ("필요가 있습니다", "필요함"),
        ("필요가 있음", "필요함"),
        ("권장합니다", "권장함"),
        ("권장됩니다", "권장됨"),
        ("검토합니다", "검토함"),
        ("검토됩니다", "검토됨"),
        ("확인합니다", "확인함"),
        ("확인됩니다", "확인됨"),
        ("진행합니다", "진행함"),
        ("진행됩니다", "진행됨"),
        ("연계합니다", "연계함"),
        ("연계됩니다", "연계됨"),
        ("계획합니다", "계획함"),
        ("수립합니다", "수립함"),
        ("안내합니다", "안내함"),
        ("관찰됩니다", "관찰됨"),
        ("나타납니다", "나타남"),
        ("보입니다", "보임"),
        ("예정입니다", "예정임"),
        ("예정입니다", "예정임"),
        ("되었습니다", "됨"),
        ("됩니다", "됨"),
        ("있습니다", "있음"),
        ("없습니다", "없음"),
        ("하였습니다", "함"),
        ("했습니다", "함"),
        ("합니다", "함"),
    ]
    for old, new in replacements:
        if text_core.endswith(old):
            text_core = text_core[: -len(old)] + new
            break
    # 너무 설명문처럼 끝나는 경우 중 일부를 회의록 문체로 정리
    text_core = text_core.replace("할 수 있습니다", "할 수 있음").replace("할 수 있음음", "할 수 있음")
    return prefix + text_core + end_punct


def nominalize_meeting_text(value: Any) -> Any:
    if isinstance(value, list):
        return [nominalize_meeting_text(v) for v in value]
    if not isinstance(value, str):
        return value
    lines = value.splitlines()
    converted = []
    for line in lines:
        parts = re.split(r"(?<=\.)\s+", line.strip()) if line.strip() else [line]
        converted.append(" ".join(_nominalize_korean_sentence(part) for part in parts if part != ""))
    return "\n".join(converted)


def normalize_meeting_record_nominal_style(data: Dict[str, Any]) -> Dict[str, Any]:
    cleaned = copy.deepcopy(data or {})
    mr = cleaned.setdefault("meeting_record", {})
    for key in ["agenda", "meeting_content", "support_plan"]:
        if key in mr:
            mr[key] = nominalize_meeting_text(mr.get(key))
    # 최종 결정사항은 회의 후 교사가 작성하도록 항상 비워 둔다.
    mr["decision_items"] = []
    return cleaned

def _set_cell_text(cell: Any, text: str) -> None:
    cell.text = str(text or "")


def _join_list(values: Any) -> str:
    if isinstance(values, list):
        return "\n".join([f"- {display_text(v)}" for v in values])
    return display_text(values or "")


def fill_meeting_docx(template_path: Path, output_path: Path, basic: Dict[str, Any], llm_result: Dict[str, Any]) -> None:
    from docx import Document
    import shutil
    llm_result = normalize_meeting_record_nominal_style(llm_result)
    shutil.copyfile(template_path, output_path)
    doc = Document(str(output_path))
    mr = llm_result.get("meeting_record", {})
    table = doc.tables[0]
    # 새 서식: 제목에 회차 없음, 지원계획/결정사항 칸 사용
    # 제목은 템플릿 서식을 그대로 유지합니다.
    table.cell(2, 1).text = basic.get("meeting_date", "")
    table.cell(2, 3).text = basic.get("writer", "")
    table.cell(2, 5).text = basic.get("place", "")
    table.cell(3, 1).text = basic.get("attendees", "")
    table.cell(4, 1).text = display_text(basic.get("agenda", "") or mr.get("agenda", ""))
    table.cell(5, 1).text = display_text(mr.get("meeting_content", ""))
    support_plan = mr.get("support_plan", None)
    if support_plan is None:
        support_plan = mr.get("decision_items", [])
    table.cell(6, 1).text = _join_list(support_plan)
    table.cell(7, 1).text = ""
    doc.save(str(output_path))

def validate_generated_docx(path: Path) -> Tuple[bool, str]:
    if not path.exists() or path.stat().st_size <= 0:
        return False, "파일이 생성되지 않았습니다."
    try:
        from docx import Document
        doc = Document(str(path))
        text = "\n".join([p.text for p in doc.paragraphs] + [cell.text for table in doc.tables for row in table.rows for cell in row.cells])
        if "{{" in text or "}}" in text:
            return False, "남은 placeholder가 있습니다."
        if re.search(r"\[[GWFISO]\d\]", text):
            return False, "서식 표시자가 남아 있습니다."
        banned = [w for w in HARD_BANNED_COMMON if w in text]
        if banned:
            return False, "금지 표현이 문서에 포함되었습니다: " + ", ".join(banned)
    except Exception as exc:
        return False, f"docx 검증 중 오류: {exc}"
    return True, "검증 통과"


def render_document_generation_section() -> None:
    if not st.session_state.get("structured_counseling_analysis") or not st.session_state.get("resource_recommendation_explanation"):
        st.info("상담 결과 분석과 지원기관 검색이 완료되면 회의록을 생성할 수 있습니다.")
        return
    st.write("회의록 기본정보")
    c1, c2, c3 = st.columns(3)
    with c1:
        meeting_date = st.text_input("일시", value=str(date.today()))
    with c2:
        writer = st.text_input("작성자")
        place = st.text_input("장소", value="교내 협의실")
    with c3:
        agenda_user = st.text_input("안건", value="대상 학생 맞춤지원 방안 협의")
    attendees = st.text_area("참석자", placeholder="예: 담임교사, 학년부장, 상담교사, 보건교사")
    if not get_gemini_api_key():
        st.warning("Gemini API 키가 설정되어 있지 않습니다. Streamlit secrets 또는 환경변수에 GEMINI_API_KEY를 설정해 주세요.")
    if st.button("회의록 생성하기", type="primary", use_container_width=True, key="btn_generate_docs"):
        if not get_gemini_api_key():
            st.warning("API 키가 설정되지 않아 회의록 초안 생성을 실행할 수 없습니다.")
        else:
            rag = st.session_state.get("rag_search_results", {})
            rec = st.session_state.get("resource_recommendation_explanation", {})
            allowed_names = [r.get("resource_name") for r in rec.get("recommended_resources", [])]
            payload = build_document_generation_payload(st.session_state.get("structured_counseling_analysis", {}), rec, rag, st.session_state.get("first_check_result", {}), "meeting_record")
            payload_hash = stable_payload_hash(payload)
            if use_cached_llm_result("generated_document_json", "generated_document_json_payload_hash", payload_hash):
                st.info("같은 입력값으로 이미 생성된 회의록 초안을 재사용합니다. API를 다시 호출하지 않았습니다.")
                result = {"success": True, "data": st.session_state["generated_document_json"], "warnings": []}
            else:
                with st.spinner("회의록 내용을 생성하고 있습니다..."):
                    result = call_llm_with_validation(
                        build_document_generation_system_prompt(),
                        build_document_generation_user_prompt(payload),
                        validate_meeting_generation_output,
                        build_document_generation_repair_prompt,
                        validation_kwargs={"allowed_resource_names": allowed_names},
                    )
            if result["success"]:
                result["data"] = normalize_meeting_record_nominal_style(result["data"])
                save_llm_result("generated_document_json", "generated_document_json_payload_hash", payload_hash, result["data"])
                out_dir = APP_DIR / "outputs"
                out_dir.mkdir(exist_ok=True)
                generated_files = {}
                meeting_basic = {"meeting_date": meeting_date, "writer": writer, "place": place, "attendees": attendees, "agenda": agenda_user}
                tpl = TEMPLATE_DIR / "회의록.docx"
                out = out_dir / f"통합지원팀_회의록_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                if not tpl.exists():
                    st.error("templates/회의록.docx 템플릿을 찾을 수 없습니다.")
                else:
                    fill_meeting_docx(tpl, out, meeting_basic, result["data"])
                    ok, msg = validate_generated_docx(out)
                    if not ok:
                        st.warning("회의록 확인 필요: " + msg)
                    generated_files["meeting"] = out
                    st.session_state["generated_docx_files"] = {k: str(v) for k, v in generated_files.items()}
                    st.success("회의록 생성이 완료되었습니다.")
                    mr = result["data"].get("meeting_record", {})
                    st.markdown("<div class='callout'>회의 후 최종 결정사항 칸은 교사가 직접 작성하도록 비워두었습니다.</div>", unsafe_allow_html=True)
                    with st.expander("생성 내용 미리보기", expanded=False):
                        st.write("안건: " + str(mr.get("agenda", "")))
                        st.write("내용: " + str(mr.get("meeting_content", "")))
                        st.write("지원계획")
                        for item in mr.get("support_plan", mr.get("decision_items", [])):
                            st.write("- " + str(item))
            else:
                st.warning(user_friendly_generation_error(result.get("error"), "회의록 생성"))
    files = st.session_state.get("generated_docx_files", {})
    if files.get("meeting") and Path(files["meeting"]).exists():
        st.download_button("회의록 DOCX 다운로드", Path(files["meeting"]).read_bytes(), Path(files["meeting"]).name, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

def render_followup_workflow_sections() -> None:
    # 각 단계는 이전 단계가 완료된 뒤에만 표시해 실제 사용자가 한 번에 과도한 정보를 보지 않도록 한다.
    with st.container(border=True):
        render_workflow_section_header("2", "2차 상담 질문 만들기", "체크리스트 결과를 바탕으로 학생과 대화할 때 활용할 질문을 제안합니다.")
        render_counseling_question_section()

    if not st.session_state.get("generated_counseling_questions"):
        return
    with st.container(border=True):
        render_workflow_section_header("3", "2차 상담 결과 정리", "상담 메모에서 확인된 신호를 간단히 정리하고, 실제로 함께 볼 지원 영역을 좁힙니다.")
        render_counseling_analysis_section()

    if not st.session_state.get("structured_counseling_analysis"):
        return
    with st.container(border=True):
        render_workflow_section_header("4", "지원기관 추천", "상담 결과와 학교 위치를 바탕으로 검토할 수 있는 지원기관을 추천합니다.")
        render_rag_search_section()

    if not st.session_state.get("resource_recommendation_explanation") and not st.session_state.get("rag_search_results"):
        return
    with st.container(border=True):
        render_workflow_section_header("5", "회의록 생성", "상담 결과와 지원기관 후보를 바탕으로 회의록 초안을 생성합니다.")
        render_document_generation_section()

def chart_bar(data: pd.DataFrame, x: str, y: str, title: str = "") -> None:
    if px is None:
        # Streamlit 기본 차트도 마우스 드래그 확대가 생길 수 있어, 가능한 경우 Plotly 경로를 우선 사용한다.
        st.bar_chart(data.set_index(x)[y])
    else:
        fig = px.bar(data, x=x, y=y, text=y, title=title)
        fig.update_layout(height=300, margin=dict(l=10, r=10, t=35, b=10), dragmode=False)
        fig.update_xaxes(fixedrange=True)
        fig.update_yaxes(fixedrange=True)
        fig.update_traces(textposition="outside")
        # 대시보드용 요약 그래프는 값 확인 목적이므로 확대/축소 인터랙션을 비활성화해
        # 사용자가 확대 후 원래 상태로 돌아가지 못하는 혼선을 줄인다.
        st.plotly_chart(
            fig,
            use_container_width=True,
            config={
                "displayModeBar": False,
                "scrollZoom": False,
                "doubleClick": False,
                "staticPlot": False,
            },
        )


def page_dashboard() -> None:
    role = st.session_state.role
    scope = f"{st.session_state.homeroom_grade} {st.session_state.homeroom_class}" if role == ROLE_HOMEROOM else "전교"
    render_page_title(
        f"{scope} 학생맞춤통합지원 대시보드",
        "담임교사는 본인 반만, 학생맞춤통합지원담당교원은 전교 학생 상태를 조회합니다.",
    )
    df = get_view_students()
    total = len(df)
    deep = int(df["최종단계"].isin(["심층 파악 필요", "심층 파악 권고"]).sum())
    watch = int((df["최종단계"] == "주의 및 탐색").sum())
    red = int(df["RedFlag"].astype(bool).sum())
    meeting = deep

    metric_cols = st.columns(5)
    metrics = [
        ("조회 학생", f"{total}명", role),
        ("심층 파악 필요/권고", f"{deep}명", "2차 상담 질문 후보"),
        ("주의 및 탐색", f"{watch}명", "담임 면담 후보"),
        ("우선 확인 신호", f"{red}명", "긴급 확인 관련"),
        ("회의자료 후보", f"{meeting}건", "학생맞춤통합지원 회의 검토"),
    ]
    for col, (label, value, help_text) in zip(metric_cols, metrics):
        with col:
            st.markdown(metric_card(label, value, help_text), unsafe_allow_html=True)

    left, right = st.columns([1.35, 1])
    with left:
        st.markdown("<div class='panel'><div class='panel-title'>지원 영역별 분포</div>", unsafe_allow_html=True)
        domain_counts = pd.DataFrame({"지원영역": [display_area_name(a) for a in SUPPORT_AREAS], "학생수": [(df[a] > 0).sum() for a in SUPPORT_AREAS]})
        chart_bar(domain_counts, "지원영역", "학생수")
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("<div class='panel'><div class='panel-title'>우선 검토 학생 목록</div>", unsafe_allow_html=True)
        order = {"심층 파악 필요": 0, "심층 파악 권고": 1, "주의 및 탐색": 2, "일상적 관찰": 3}
        show = df[df["최종단계"].isin(["심층 파악 필요", "심층 파악 권고", "주의 및 탐색"])].copy()
        show["정렬"] = show["최종단계"].map(order)
        show["합계"] = show[SUPPORT_AREAS].sum(axis=1)
        show = show.sort_values(["정렬", "RedFlag", "합계"], ascending=[True, False, False])
        if show.empty:
            st.info("현재 조회 범위에서 우선 검토 학생이 없습니다.")
        for _, row in show.head(20).iterrows():
            render_student_card(row)
        st.markdown("</div>", unsafe_allow_html=True)

    with right:
        st.markdown("<div class='panel'><div class='panel-title'>행동 단계 분포</div>", unsafe_allow_html=True)
        stage_counts = df["최종단계"].value_counts().reindex(STATUS_ORDER, fill_value=0).reset_index()
        stage_counts.columns = ["단계", "학생수"]
        chart_bar(stage_counts, "단계", "학생수")
        st.markdown("</div>", unsafe_allow_html=True)


# -----------------------------------------------------------------------------
# 페이지: 1차 체크리스트
# -----------------------------------------------------------------------------
def render_checklist_input(items_df: pd.DataFrame, selected_student: str) -> Dict[str, int]:
    responses: Dict[str, int] = {}
    if items_df.empty:
        st.error("1차 체크리스트 CSV를 불러올 수 없습니다.")
        return responses

    score_labels = {
        0: "관찰되지 않음",
        1: "약하게/가끔 관찰됨",
        2: "뚜렷하게/자주 관찰됨",
    }
    grouped = items_df.groupby(["domain_order", "domain_label"], dropna=False, sort=True)
    for (_, domain_label), sub in grouped:
        with st.expander(display_text(domain_label), expanded=True):
            for _, row in sub.iterrows():
                item_id = normalize_text(row.get("item_id"))
                item_text = clean_checklist_item_label(row.get("item_text"))
                default_value = st.session_state.checklist_responses.get(selected_student, {}).get(item_id, 0)
                value = st.radio(
                    item_text,
                    [0, 1, 2],
                    index=[0, 1, 2].index(int(default_value)),
                    format_func=lambda x: score_labels[x],
                    horizontal=True,
                    key=f"score_{selected_student}_{item_id}",
                )
                responses[item_id] = int(value)
    return responses



def page_first_checklist() -> None:
    render_page_title(
        "1차 체크리스트 입력",
        "교사가 관찰한 신호를 입력하면 학생맞춤통합지원 검토 단계와 다음 조치를 안내합니다.",
    )

    items_df = get_active_items_df()
    rule_map_df = st.session_state.get("rule_map_df")
    deep_rules_df = st.session_state.get("deep_rules_df")

    df = get_view_students()
    if df.empty:
        st.warning("현재 조회 범위에 학생이 없습니다.")
        return

    student_options = [f"{row['학생코드']} | {row['이름']} | {row['학년']} {row['반']}" for _, row in df.iterrows()]
    preferred_student = st.session_state.get("selected_student_for_checklist")
    default_index = 0
    if preferred_student:
        for i, label in enumerate(student_options):
            if label.startswith(f"{preferred_student} |"):
                default_index = i
                break
    selected_label = st.selectbox("학생 선택", student_options, index=default_index)
    selected_student = selected_label.split(" | ")[0]
    st.session_state["selected_student_for_checklist"] = selected_student

    render_workflow_section_header("1", "1차 체크리스트", "학생의 학교생활에서 관찰된 신호를 입력합니다.")
    responses = render_checklist_input(items_df, selected_student)
    calc_clicked = st.button("체크리스트 결과 계산", type="primary", use_container_width=True)

    if not responses:
        return

    if checklist_input_changed(selected_student, responses) and not calc_clicked:
        clear_pipeline_from("checklist")
        st.info("체크리스트 입력이 변경되었습니다. 다시 계산하면 이후 단계가 새 결과에 맞춰 열립니다.")
        return

    if calc_clicked:
        calculate_and_store_checklist_for_student(
            selected_student,
            responses,
            items_df,
            rule_map_df,
            deep_rules_df,
            update_student_row=True,
        )

    if not st.session_state.get("first_check_result") or st.session_state.get("last_checklist_student") != selected_student:
        st.info("체크리스트 입력 후 ‘체크리스트 결과 계산’을 누르면 결과가 표시됩니다.")
        return

    first_check_result = st.session_state.get("first_check_result", {})
    context_result = st.session_state.get("context_result", {})
    red_flag_result = st.session_state.get("red_flag_result", {})
    counseling_areas = st.session_state.get("counseling_consideration_areas", [])

    active_deep_rules = st.session_state.get("active_deep_rules", [])
    direct_areas, related_areas = get_checklist_direct_and_related_areas(first_check_result, counseling_areas, active_deep_rules)
    direct_areas_text = ", ".join(direct_areas) if direct_areas else "현재 뚜렷한 우선 영역 없음"
    related_areas_help = area_help_text("심층 분석상 함께 고려할 수 있는 영역", related_areas, "심층 분석상 추가 고려 영역 없음")

    final_stage = first_check_result.get("final_action_stage") or context_result.get("final_action_stage", "-")
    final_reason = first_check_result.get("final_action_reason") or context_result.get("final_action_reason", "")
    action = "2차 상담 질문 생성" if first_check_result.get("activate_counseling_form") else "담임교사 면담 및 추가 관찰"

    st.markdown("<div class='panel'><div class='panel-title'>체크리스트 결과</div>", unsafe_allow_html=True)
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.markdown(metric_card("안내 단계", str(final_stage), "교사 검토용 안내"), unsafe_allow_html=True)
    with c2:
        st.markdown(metric_card("주요 지원 영역", direct_areas_text, related_areas_help), unsafe_allow_html=True)
    with c3:
        st.markdown(metric_card("우선 확인 신호", "있음" if red_flag_result.get("urgent_flag") else "없음", "필요 시 먼저 확인"), unsafe_allow_html=True)
    with c4:
        st.markdown(metric_card("권장 조치", action, "교사 참고"), unsafe_allow_html=True)
    if final_reason:
        st.info(final_reason)
    if final_stage == "일상적 관찰":
        st.caption("현재는 평소 관찰과 라포 형성을 유지하는 단계입니다.")
    elif final_stage == "주의 및 탐색":
        st.caption("담임교사의 가벼운 면담으로 최근 변화와 어려움을 탐색할 수 있습니다.")
    elif final_stage == "심층 파악 권고":
        st.caption("2차 상담 질문을 통해 한 번 더 확인하는 것이 권장됩니다.")
    else:
        st.caption("학생과의 추가 상담을 통해 어려움을 조금 더 구체적으로 확인할 수 있습니다.")
    st.markdown("</div>", unsafe_allow_html=True)

    render_followup_workflow_sections()

# -----------------------------------------------------------------------------
# 페이지: 학생 상세 / 지원 현황 / 데이터 연결 안내
# -----------------------------------------------------------------------------
def _get_generated_meeting_file() -> Optional[Path]:
    files = st.session_state.get("generated_docx_files", {}) or {}
    meeting_path = files.get("meeting")
    if meeting_path and Path(meeting_path).exists():
        return Path(meeting_path)
    return None


def render_student_detail_followup_summary(selected_student: str, student: pd.Series) -> None:
    """학생 상세 리포트의 오른쪽 영역에 현재까지 완료된 후속 결과를 간단히 표시한다."""
    areas = [display_area_name(a) for a in SUPPORT_AREAS if int(student.get(a, 0)) > 0]
    st.write("우선 검토 영역: " + (", ".join(areas) if areas else "현재 뚜렷한 우선 영역 없음"))
    st.write("관찰 신호: " + display_text(student.get("주요신호")))

    # 현재 세션에서 선택 학생에 대해 2차 상담 결과가 생성된 경우에는
    # 1차 체크리스트 안내 문구 대신 실제 후속 결과를 보여준다.
    same_student = st.session_state.get("last_checklist_student") == selected_student
    analysis = st.session_state.get("structured_counseling_analysis") if same_student else None
    meeting_file = _get_generated_meeting_file() if same_student else None

    if analysis:
        analysis = postprocess_counseling_analysis_result(analysis, st.session_state.get("teacher_counseling_note", ""))
        summary = analysis.get("analysis_summary", {})
        one_line = summary.get("support_needed_reason") or summary.get("one_sentence_summary") or "상담 결과를 바탕으로 지원 검토 방향을 정리했습니다."
        status = summary.get("support_needed") or "지원 검토 필요"
        target_areas = derive_integrated_support_areas(analysis)
        area_text = ", ".join(display_area_list(target_areas)) if target_areas else display_area_name(analysis.get("primary_area", "-"))

        st.markdown(
            f"""
            <div class='soft-card compact-card'>
                <div class='subtle-label'>2차 상담 결과 정리</div>
                <div class='mini-title'>{display_text(status)}</div>
                <div class='readable-note'>{display_text(one_line)}</div>
                <div class='subtle-label' style='margin-top:14px;'>상담에서 확인된 지원 영역</div>
                <div class='area-inline'>{area_text}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        rec = st.session_state.get("resource_recommendation_explanation") or {}
        rec_count = len(rec.get("recommended_resources", []) or [])
        if rec_count:
            st.caption(f"지원기관 후보 {rec_count}곳을 추천했습니다.")
        if meeting_file:
            st.download_button(
                "회의록 DOCX 다운로드",
                meeting_file.read_bytes(),
                meeting_file.name,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key=f"student_detail_meeting_download_{selected_student}",
            )
        else:
            st.info("회의록을 생성하면 이곳에서 바로 다운로드할 수 있습니다.")
        return

    if student["최종단계"] == "심층 파악 필요":
        st.error("심층 파악이 필요한 상태입니다. 1차 체크리스트 결과를 바탕으로 2차 상담 질문 생성을 검토합니다.")
        if st.button("1차 체크리스트로 이동해 상담 질문 만들기", type="primary", use_container_width=True, key=f"open_checklist_{selected_student}"):
            prepare_student_checklist_and_open(selected_student, student)
            st.rerun()
    elif student["최종단계"] == "주의 및 탐색":
        st.warning("담임교사의 가벼운 면담과 관찰 기록 업데이트를 권장합니다.")
        if st.button("1차 체크리스트에서 확인하기", use_container_width=True, key=f"open_checklist_watch_{selected_student}"):
            prepare_student_checklist_and_open(selected_student, student)
            st.rerun()
    else:
        st.info("현재는 일상적 관찰 단계입니다.")


def page_student_detail() -> None:
    render_page_title("학생 상세 리포트", "선택 학생의 현재 지원 신호와 다음 조치 후보를 확인합니다.")
    df = get_view_students()
    if df.empty:
        st.warning("현재 조회 범위에 학생이 없습니다.")
        return
    selected = st.selectbox("학생 선택", df["학생코드"].tolist())
    student = df[df["학생코드"] == selected].iloc[0]

    col1, col2 = st.columns([1, 1.4])
    with col1:
        render_student_card(student)
        st.markdown("<div class='panel'><div class='panel-title'>학생 요약</div>", unsafe_allow_html=True)
        table = pd.DataFrame(
            [
                ["학생코드", student["학생코드"]],
                ["학년·반", f"{student['학년']} {student['반']}"],
                ["최종 단계", student["최종단계"]],
                ["우선 확인 신호", "있음" if student["RedFlag"] else "없음"],
                ["담당자", student["담당자"]],
                ["기한", student["기한"]],
            ],
            columns=["항목", "내용"],
        )
        st.dataframe(table, use_container_width=True, hide_index=True)
        st.markdown("</div>", unsafe_allow_html=True)

    with col2:
        st.markdown("<div class='panel'><div class='panel-title'>지원 검토 메모</div>", unsafe_allow_html=True)
        render_student_detail_followup_summary(selected, student)
        st.markdown("</div>", unsafe_allow_html=True)


def make_status_table(df: pd.DataFrame) -> pd.DataFrame:
    table = df.copy()
    table["지원영역"] = table.apply(lambda r: ", ".join([display_area_name(a) for a in SUPPORT_AREAS if int(r[a]) > 0]) or "-", axis=1)
    table["주요신호"] = table["주요신호"].map(display_text)
    table["다음 할 일"] = table["권장Action"].map(display_text)
    return table[["학생코드", "학년", "반", "최종단계", "RedFlag", "지원영역", "주요신호", "다음 할 일", "담당자", "기한"]]


def to_excel_bytes(df: pd.DataFrame) -> bytes:
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="지원현황")
    return output.getvalue()


def page_status_table() -> None:
    render_page_title("지원 현황표", "조회 권한에 따른 지원 현황을 표로 보고 CSV/Excel로 다운로드합니다.")
    df = get_view_students()
    stage_filter = st.multiselect("단계", STATUS_ORDER, default=["심층 파악 필요", "주의 및 탐색"])
    if stage_filter:
        df = df[df["최종단계"].isin(stage_filter)]
    table = make_status_table(df)
    st.dataframe(table, use_container_width=True, hide_index=True)

    col1, col2 = st.columns(2)
    with col1:
        st.download_button("CSV 다운로드", data=table.to_csv(index=False).encode("utf-8-sig"), file_name="학생맞춤통합지원_지원현황표.csv", mime="text/csv", use_container_width=True)
    with col2:
        st.download_button("Excel 다운로드", data=to_excel_bytes(table), file_name="학생맞춤통합지원_지원현황표.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)


# -----------------------------------------------------------------------------
# 초기화 / 사이드바
# -----------------------------------------------------------------------------
def init_state() -> None:
    if "role" not in st.session_state:
        st.session_state.role = ROLE_HOMEROOM
    if "homeroom_grade" not in st.session_state:
        st.session_state.homeroom_grade = DEFAULT_GRADE
    if "homeroom_class" not in st.session_state:
        st.session_state.homeroom_class = DEFAULT_CLASS
    if "checklist_responses" not in st.session_state:
        st.session_state.checklist_responses = {}
    if "current_page" not in st.session_state:
        st.session_state.current_page = "교사 대시보드"
    if "selected_student_for_checklist" not in st.session_state:
        st.session_state.selected_student_for_checklist = None

    school_db = load_school_databases()
    st.session_state.school_db = school_db
    for err in school_db.get("errors", []):
        st.sidebar.error(err)

    school_info_df = school_db["school_info_df"]
    if "selected_school_name" not in st.session_state:
        if "공항고등학교" in school_info_df["학교명"].astype(str).tolist():
            st.session_state.selected_school_name = "공항고등학교"
        else:
            st.session_state.selected_school_name = str(school_info_df.iloc[0]["학교명"])

    selected_info_row = find_matching_row(school_info_df, st.session_state.selected_school_name)
    if selected_info_row is None:
        selected_info_row = school_info_df.iloc[0]
        st.session_state.selected_school_name = str(selected_info_row["학교명"])
    st.session_state.selected_school_info = school_row_to_dict(selected_info_row)

    district = st.session_state.selected_school_info.get("자치구", "")
    st.session_state.selected_school_context = find_matching_row(
        school_db.get("school_context_df"),
        st.session_state.selected_school_name,
        district,
    )
    st.session_state.selected_region_context = find_region_row(school_db.get("region_context_df"), district)

    # 체크리스트/규칙 CSV
    items_df, items_path, items_err = load_optional_csv(CSV_PATHS["checklist"], "1차 체크리스트")
    deep_rules_df, deep_rules_path, deep_rules_err = load_optional_csv(CSV_PATHS["deep_rules"], "심층 유도 규칙")
    rule_map_df, rule_map_path, rule_map_err = load_optional_csv(CSV_PATHS["rule_map"], "체크리스트-규칙 매핑")
    for err in [items_err, deep_rules_err, rule_map_err]:
        if err:
            st.sidebar.error(err)
    st.session_state.checklist_items_df = items_df if items_df is not None else pd.DataFrame()
    st.session_state.deep_rules_df = deep_rules_df if deep_rules_df is not None else pd.DataFrame()
    st.session_state.rule_map_df = rule_map_df if rule_map_df is not None else pd.DataFrame()
    st.session_state.checklist_path_name = items_path.name if items_path else "없음"
    st.session_state.deep_rules_path_name = deep_rules_path.name if deep_rules_path else "없음"
    st.session_state.rule_map_path_name = rule_map_path.name if rule_map_path else "없음"

    if "students_school_name" not in st.session_state or st.session_state.students_school_name != st.session_state.selected_school_name:
        st.session_state.students = generate_demo_students(st.session_state.selected_school_name)
        st.session_state.students_school_name = st.session_state.selected_school_name


def render_sidebar() -> str:
    st.sidebar.markdown("### 학교 선택")
    school_info_df = st.session_state.school_db["school_info_df"]
    school_names = school_info_df["학교명"].astype(str).drop_duplicates().tolist()
    if st.session_state.selected_school_name not in school_names:
        st.session_state.selected_school_name = school_names[0]
    selected_school = st.sidebar.selectbox(
        "학교명",
        school_names,
        index=school_names.index(st.session_state.selected_school_name),
        help="선택한 학교의 기본 정보를 불러옵니다.",
    )
    if selected_school != st.session_state.selected_school_name:
        st.session_state.selected_school_name = selected_school
        st.rerun()

    st.sidebar.markdown("### 담임 반 설정")
    st.session_state.homeroom_grade = st.sidebar.selectbox("학년", ["1학년", "2학년", "3학년"], index=["1학년", "2학년", "3학년"].index(st.session_state.homeroom_grade))
    st.session_state.homeroom_class = st.sidebar.selectbox("반", ["1반", "2반", "3반", "4반"], index=["1반", "2반", "3반", "4반"].index(st.session_state.homeroom_class))

    st.sidebar.divider()
    st.sidebar.markdown("### 메뉴")
    menu_options = ["교사 대시보드", "1차 체크리스트", "학생 상세 리포트", "지원 현황표"]
    pending_page = st.session_state.pop("pending_page", None)
    if pending_page in menu_options:
        st.session_state.current_page = pending_page
    if st.session_state.get("current_page") not in menu_options:
        st.session_state.current_page = menu_options[0]
    page = st.sidebar.radio(
        "화면 선택",
        menu_options,
        index=menu_options.index(st.session_state.current_page),
        label_visibility="collapsed",
        key="current_page",
    )

    school = st.session_state.selected_school_info
    st.sidebar.divider()
    st.sidebar.markdown("### 현재 학교 정보")
    st.sidebar.write(f"학교: {school.get('학교명')}")
    st.sidebar.write(f"학교급: {school.get('학교급')}")
    st.sidebar.write(f"자치구: {school.get('자치구')}")
    st.sidebar.write(f"교육지원청: {school.get('교육지원청')}")
    st.sidebar.write(f"학생 수: {int(school.get('학교_학생수', 0)):,}명" if school.get('학교_학생수', 0) else "학생 수: -")
    st.sidebar.write(f"교원 1인당 학생수: {school.get('학교_교원1인당학생수', 0):.1f}명" if school.get('학교_교원1인당학생수', 0) else "교원 1인당 학생수: -")
    st.sidebar.write(f"위클래스: {'있음' if school.get('위클래스_있음') else '없음'}")
    st.sidebar.write(f"전문상담교사: {school.get('전문상담교사_수')}명")
    st.sidebar.write(f"보건교사: {school.get('보건교사_수')}명")
    st.sidebar.write(f"진로상담실: {'있음' if school.get('진로상담실_있음') else '없음'}")
    return page

# -----------------------------------------------------------------------------
# 메인
# -----------------------------------------------------------------------------
def main() -> None:
    inject_css()
    init_state()
    render_header()
    render_role_switch()
    page = render_sidebar()

    if page == "교사 대시보드":
        page_dashboard()
    elif page == "1차 체크리스트":
        page_first_checklist()
    elif page == "학생 상세 리포트":
        page_student_detail()
    elif page == "지원 현황표":
        page_status_table()

    st.markdown(
        """
        <div class="footer-note">
        본 화면은 학생맞춤통합지원 서비스의 시연용 화면입니다. 학생 지원 여부를 자동 확정하지 않으며,
        담임교사와 학교 협의체의 검토를 돕는 참고자료로 설계되었습니다.
        </div>
        """,
        unsafe_allow_html=True,
    )


if __name__ == "__main__":
    main()
